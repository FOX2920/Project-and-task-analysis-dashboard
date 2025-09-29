# Standard library imports
import io
import json
import os
import re
import time
from datetime import datetime, timedelta
from collections import defaultdict, Counter
from typing import Dict, List, Tuple, Optional, Any, Union

# Third-party imports
import numpy as np
import pandas as pd
import requests
from bs4 import BeautifulSoup

# Streamlit and visualization
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import seaborn as sns
import matplotlib.pyplot as plt

# Document processing
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Optional AI imports
try:
    from google import genai
    from google.genai import types
    GENAI_AVAILABLE = True
except ImportError:
    GENAI_AVAILABLE = False
    genai = None
    types = None
WEWORK_ACCESS_TOKEN = os.getenv('BASE_API_KEY')
ACCOUNT_ACCESS_TOKEN = os.getenv('ACCOUNT_ACCESS_TOKEN')
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')
# Note: AI availability warning will be shown in main() function

# Configuration Constants
class Config:
    """Application configuration constants"""

    # Network and Performance
    REQUEST_TIMEOUT = 30
    MAX_DISPLAY_ITEMS = 15
    MAX_EXPORT_SIZE = 50 * 1024 * 1024  # 50MB


    # Analysis Options
    ANALYSIS_MODES = ["Cơ bản", "AI", "Cả hai"]
    FEATURE_MODES = ["Phân tích theo thời gian", "Employee 360", "Phân tích Dự án", "Cả ba"]
    TIME_FILTERS = [
        "Tháng hiện tại", "Quý hiện tại", "Năm hiện tại",
        "3 tháng gần nhất", "6 tháng gần nhất", "Tất cả"
    ]

    # Risk Assessment
    RISK_LEVELS = {
        'VERY_HIGH': {'score': 6, 'label': 'Rủi ro rất cao', 'color': 'red'},
        'HIGH': {'score': 4, 'label': 'Rủi ro cao', 'color': 'orange'},
        'MEDIUM': {'score': 2, 'label': 'Rủi ro trung bình', 'color': 'yellow'},
        'LOW': {'score': 0, 'label': 'Rủi ro thấp', 'color': 'green'}
    }

    # Quality Grades
    QUALITY_GRADES = [
        (90, "Xuất sắc (A+)"),
        (80, "Tốt (A)"),
        (70, "Khá (B+)"),
        (60, "Trung bình khá (B)"),
        (50, "Trung bình (C)"),
        (0, "Cần cải thiện (D)")
    ]

    # UI Constants
    COLORS = {
        'success': '#2ECC71',
        'warning': '#F39C12',
        'error': '#E74C3C',
        'info': '#3498DB',
        'primary': '#9B59B6'
    }

    # Chart Configuration
    CHART_HEIGHTS = {
        'dashboard': 1000,
        'timeline': 600,
        'radar': 500,
        'bar': 400
    }

# Utility Functions
class Utils:
    """Utility functions for data processing and validation"""

    @staticmethod
    def safe_timestamp_convert(timestamp: Any) -> Optional[str]:
        """Safely convert timestamp to datetime string"""
        try:
            if timestamp and str(timestamp).strip() and int(timestamp) != 0:
                return datetime.fromtimestamp(int(timestamp)).strftime('%Y-%m-%d')
        except (ValueError, TypeError, OverflowError):
            pass
        return None

    @staticmethod
    def clean_html_content(content: Any) -> str:
        """Clean HTML content and extract plain text"""
        if not isinstance(content, str):
            return ""

        soup = BeautifulSoup(content, 'html.parser')

        # Remove style attributes
        for tag in soup.find_all(style=True):
            del tag['style']

        text = str(soup)

        # Replace HTML tags with readable text
        replacements = {
            '<p>': '', '</p>': '\n',
            '<ul>': '', '</ul>': '',
            '<li>': '- ', '</li>': '\n',
            '<br>': '\n', '<br/>': '\n',
            '<ol>': '', '</ol>': '',
            '<span>': '', '</span>': '',
            '<strong>': '', '</strong>': '',
            '&nbsp;': ' '
        }

        for old, new in replacements.items():
            text = text.replace(old, new)

        # Remove remaining HTML tags
        text = re.sub('<[^<]+?>', '', text)
        text = ' '.join(text.split())
        return text.strip()

    @staticmethod
    def calculate_completion_rate(completed: int, total: int) -> float:
        """Calculate completion rate safely"""
        return (completed / total * 100) if total > 0 else 0.0

    @staticmethod
    def get_risk_level(score: int) -> str:
        """Get risk level description from score"""
        for level_key, level_info in Config.RISK_LEVELS.items():
            if score >= level_info['score']:
                return level_info['label']
        return Config.RISK_LEVELS['LOW']['label']

    @staticmethod
    def get_quality_grade(score: float) -> str:
        """Get quality grade from score"""
        for threshold, grade in Config.QUALITY_GRADES:
            if score >= threshold:
                return grade
        return Config.QUALITY_GRADES[-1][1]

    @staticmethod
    def format_file_size(size_bytes: int) -> str:
        """Format file size in human readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"

    @staticmethod
    def validate_project_data(project_data: Dict) -> bool:
        """Validate project data structure"""
        required_fields = ['project', 'tasks']
        return all(field in project_data for field in required_fields)

    @staticmethod
    def chunk_list(items: List, chunk_size: int) -> List[List]:
        """Split list into chunks"""
        return [items[i:i + chunk_size] for i in range(0, len(items), chunk_size)]

    @staticmethod
    def safe_json_loads(data: str) -> Optional[Dict]:
        """Safely parse JSON data"""
        try:
            return json.loads(data)
        except (json.JSONDecodeError, TypeError):
            return None

    @staticmethod
    def validate_api_response(response: requests.Response) -> bool:
        """Validate API response"""
        try:
            response.raise_for_status()
            return True
        except requests.exceptions.RequestException:
            return False

    @staticmethod
    def get_safe_dict_value(data: Dict, key: str, default: Any = None) -> Any:
        """Safely get value from dictionary"""
        return data.get(key, default) if isinstance(data, dict) else default

    @staticmethod
    def calculate_percentage(numerator: int, denominator: int) -> float:
        """Calculate percentage safely"""
        return Utils.calculate_completion_rate(numerator, denominator)

    @staticmethod
    def truncate_text(text: str, max_length: int = 50) -> str:
        """Truncate text to max length with ellipsis"""
        if len(text) <= max_length:
            return text
        return text[:max_length - 3] + "..."

    @staticmethod
    def get_nested_value(data: Dict, keys: List[str], default: Any = None) -> Any:
        """Get nested dictionary value safely"""
        current = data
        for key in keys:
            if not isinstance(current, dict) or key not in current:
                return default
            current = current[key]
        return current

# Decorators for better error handling
def handle_api_errors(func):
    """Decorator to handle API errors gracefully"""
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except requests.exceptions.RequestException as e:
            st.error(f"❌ Lỗi kết nối API: {str(e)}")
            return None
        except Exception as e:
            st.error(f"❌ Lỗi không xác định: {str(e)}")
            return None
    return wrapper

def validate_input(func):
    """Decorator to validate function inputs"""
    def wrapper(*args, **kwargs):
        # Add validation logic here
        return func(*args, **kwargs)
    return wrapper

def log_execution_time(func):
    """Decorator to log function execution time"""
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        execution_time = end_time - start_time
        if execution_time > 1.0:  # Log only slow operations
            print(f"Function {func.__name__} took {execution_time:.2f} seconds")
        return result
    return wrapper

# Performance and Caching Utilities
class CacheManager:
    """Simple cache manager for API responses and processed data"""

    def __init__(self, max_size: int = 100, ttl_seconds: int = 300):
        self.max_size = max_size
        self.ttl_seconds = ttl_seconds
        self._cache = {}
        self._access_times = {}

    def get(self, key: str) -> Any:
        """Get cached value if not expired"""
        if key not in self._cache:
            return None

        # Check if expired
        if time.time() - self._access_times[key] > self.ttl_seconds:
            del self._cache[key]
            del self._access_times[key]
            return None

        return self._cache[key]

    def set(self, key: str, value: Any) -> None:
        """Set cached value with cleanup if needed"""
        # Cleanup expired items
        self._cleanup_expired()

        # Remove oldest if cache is full
        if len(self._cache) >= self.max_size:
            oldest_key = min(self._access_times, key=self._access_times.get)
            if oldest_key:
                del self._cache[oldest_key]
                del self._access_times[oldest_key]

        self._cache[key] = value
        self._access_times[key] = time.time()

    def _cleanup_expired(self) -> None:
        """Remove expired cache entries"""
        current_time = time.time()
        expired_keys = [
            key for key, access_time in self._access_times.items()
            if current_time - access_time > self.ttl_seconds
        ]

        for key in expired_keys:
            del self._cache[key]
            del self._access_times[key]

    def clear(self) -> None:
        """Clear all cache"""
        self._cache.clear()
        self._access_times.clear()

class MemoryManager:
    """Memory usage monitoring and optimization"""

    @staticmethod
    def get_memory_usage() -> float:
        """Get current memory usage in MB"""
        try:
            import psutil
            process = psutil.Process()
            return process.memory_info().rss / 1024 / 1024
        except ImportError:
            return 0.0

    @staticmethod
    def optimize_dataframe(df: pd.DataFrame) -> pd.DataFrame:
        """Optimize DataFrame memory usage"""
        # Convert object columns to category where appropriate
        for col in df.columns:
            if df[col].dtype == 'object':
                # Check if column has few unique values
                if df[col].nunique() / len(df) < 0.5:
                    df[col] = df[col].astype('category')

        return df

    @staticmethod
    def chunked_processing(items: List[Any], chunk_size: int = 1000) -> List[Any]:
        """Process items in chunks to avoid memory issues"""
        results = []
        for i in range(0, len(items), chunk_size):
            chunk = items[i:i + chunk_size]
            # Process chunk here
            results.extend(chunk)
        return results

# Global cache instance
cache_manager = CacheManager()

# Data Processing Classes
class DataValidator:
    """Data validation and processing utilities"""

    @staticmethod
    def validate_employee_data(employee_data: Dict) -> bool:
        """Validate employee data structure"""
        required_fields = ['name', 'username']
        return all(field in employee_data and employee_data[field] for field in required_fields)

    @staticmethod
    def validate_project_data(project_data: Dict) -> bool:
        """Validate project data structure"""
        return Utils.validate_project_data(project_data)

    @staticmethod
    def validate_task_data(task_data: Dict) -> bool:
        """Validate task data structure"""
        required_fields = ['name', 'content']
        return all(field in task_data for field in required_fields)

    @staticmethod
    def clean_employee_list(employees: List[Dict]) -> List[Dict]:
        """Clean and filter employee list"""
        cleaned = []
        for emp in employees:
            if DataValidator.validate_employee_data(emp):
                cleaned.append({
                    'id': Utils.get_safe_dict_value(emp, 'id', ''),
                    'name': Utils.get_safe_dict_value(emp, 'name', '').strip(),
                    'username': Utils.get_safe_dict_value(emp, 'username', '').strip(),
                    'job': Utils.get_safe_dict_value(emp, 'title', '').strip(),
                    'email': Utils.get_safe_dict_value(emp, 'email', '').strip()
                })
        return cleaned

    @staticmethod
    def clean_task_list(tasks: List[Dict]) -> List[Dict]:
        """Clean and filter task list"""
        cleaned = []
        for task in tasks:
            if DataValidator.validate_task_data(task):
                cleaned.append({
                    'name': Utils.get_safe_dict_value(task, 'name', '').strip(),
                    'content': Utils.clean_html_content(Utils.get_safe_dict_value(task, 'content', '')),
                    'complete': Utils.get_safe_dict_value(task, 'complete', '0'),
                    'start_time': Utils.get_safe_dict_value(task, 'start_time', ''),
                    'deadline': Utils.get_safe_dict_value(task, 'deadline', ''),
                    'completed_time': Utils.get_safe_dict_value(task, 'completed_time', ''),
                    'username': Utils.get_safe_dict_value(task, 'username', '').strip(),
                    'tasklist': Utils.get_safe_dict_value(task, 'tasklist', {}),
                    'result': Utils.get_safe_dict_value(task, 'result', {}),
                    'followers': Utils.get_safe_dict_value(task, 'followers', [])
                })
        return cleaned

class DataProcessor:
    """Data processing utilities"""

    @staticmethod
    def group_tasks_by_project(tasks: List[Dict]) -> Dict[str, List[Dict]]:
        """Group tasks by project name"""
        grouped = {}
        for task in tasks:
            project_name = Utils.get_nested_value(task, ['project', 'name'], 'Unknown Project')
            if project_name not in grouped:
                grouped[project_name] = []
            grouped[project_name].append(task)
        return grouped

    @staticmethod
    def calculate_project_metrics(tasks: List[Dict]) -> Dict[str, Any]:
        """Calculate project metrics from tasks"""
        if not tasks:
            return {
                'total_tasks': 0,
                'completed_tasks': 0,
                'completion_rate': 0.0,
                'avg_completion': 0.0
            }

        total_tasks = len(tasks)
        completed_tasks = sum(1 for task in tasks if Utils.get_safe_dict_value(task, 'complete') == '100.00')

        completion_percentages = [
            float(Utils.get_safe_dict_value(task, 'complete', '0'))
            for task in tasks
        ]

        return {
            'total_tasks': total_tasks,
            'completed_tasks': completed_tasks,
            'completion_rate': Utils.calculate_percentage(completed_tasks, total_tasks),
            'avg_completion': sum(completion_percentages) / len(completion_percentages) if completion_percentages else 0.0
        }

    @staticmethod
    def filter_tasks_by_time_range(tasks: List[Dict], start_date: datetime, end_date: datetime) -> List[Dict]:
        """Filter tasks by time range"""
        filtered = []
        for task in tasks:
            task_start = Utils.safe_timestamp_convert(Utils.get_safe_dict_value(task, 'start_time'))
            task_complete = Utils.safe_timestamp_convert(Utils.get_safe_dict_value(task, 'completed_time'))

            if task_start and start_date <= datetime.strptime(task_start, '%Y-%m-%d') <= end_date:
                filtered.append(task)
            elif task_complete and start_date <= datetime.strptime(task_complete, '%Y-%m-%d') <= end_date:
                filtered.append(task)

        return filtered

class APIClient:
    """
    Client for handling API requests to WeWork and Account services.

    This class provides methods to fetch employee and project data from
    the WeWork API with proper error handling and data validation.
    """

    def __init__(self, goal_token: Optional[str] = None, account_token: Optional[str] = None) -> None:
        self.goal_token = goal_token or Config.WEWORK_ACCESS_TOKEN
        self.account_token = account_token or Config.ACCOUNT_ACCESS_TOKEN
        self.cache = cache_manager  # Use global cache instance
    
    def _make_request(self, url: str, data: Dict[str, Any], description: str = "") -> requests.Response:
        """
        Make HTTP request with comprehensive error handling.

        Args:
            url: API endpoint URL
            data: Request payload data
            description: Description for error logging

        Returns:
            Response object from requests

        Raises:
            requests.exceptions.RequestException: If request fails
        """
        try:
            response = requests.post(url, data=data, timeout=Config.REQUEST_TIMEOUT)
            if not Utils.validate_api_response(response):
                raise requests.exceptions.RequestException(f"API request failed: {response.status_code}")
            return response
        except requests.exceptions.RequestException as e:
            error_msg = f"Error {description}: {e}" if description else f"API request failed: {e}"
            st.error(error_msg)
            raise

    @handle_api_errors
    def get_filtered_members(self) -> pd.DataFrame:
        """
        Get filtered members from account API with caching.

        Returns:
            DataFrame containing filtered employee data
        """
        cache_key = f"members_{self.account_token}"

        # Check cache first
        cached_data = self.cache.get(cache_key)
        if cached_data is not None:
            return cached_data

        url = "https://account.base.vn/extapi/v1/group/get"
        data = {"access_token": self.account_token, "path": "aplus"}

        response = self._make_request(url, data, "fetching account members")
        response_data = Utils.safe_json_loads(response.text)

        if not response_data:
            st.error("❌ Không thể parse dữ liệu từ API")
            return pd.DataFrame()

        members = Utils.get_nested_value(response_data, ['group', 'members'], [])

        df = pd.DataFrame([
            {
                'id': str(Utils.get_safe_dict_value(m, 'id', '')),
                'name': Utils.get_safe_dict_value(m, 'name', ''),
                'username': Utils.get_safe_dict_value(m, 'username', ''),
                'job': Utils.get_safe_dict_value(m, 'title', ''),
                'email': Utils.get_safe_dict_value(m, 'email', '')
            }
            for m in members
        ])

        # Optimize DataFrame memory usage
        df = MemoryManager.optimize_dataframe(df)

        filtered_df = self._apply_member_filters(df)

        # Cache the result
        self.cache.set(cache_key, filtered_df)

        return filtered_df
    
    def _apply_member_filters(self, df: pd.DataFrame) -> pd.DataFrame:
        """Apply filters to member dataframe"""
        excluded_jobs = 'kcs|agile|khu vực|sa ti co|trainer|specialist|no|chuyên gia|xnk|vat|trưởng phòng thị trường'
        filtered_df = df[~df['job'].str.lower().str.contains(excluded_jobs, na=False)]
        # Loại bỏ các username không mong muốn
        filtered_df = filtered_df[filtered_df['username'] != 'ThuAn']
        filtered_df = filtered_df[filtered_df['username'] != 'HienNguyen']
        return filtered_df
    
    def get_employee_projects_by_time(self, username: str, time_filter: str) -> List[Dict]:
        """Get all projects where employee participated in specified time period"""
        try:
            # Calculate time range based on filter
            now = datetime.now()
            start_date = self._calculate_start_date(now, time_filter)
            
            # Get all projects
            projects_url = "https://wework.base.vn/extapi/v3/project/list"
            projects_response = self._make_request(
                projects_url, 
                {'access_token': self.goal_token},
                "fetching projects"
            )
            projects = projects_response.json().get('projects', [])
            
            employee_projects = []
            
            for project in projects:
                # Get project details
                project_detail_url = "https://wework.base.vn/extapi/v3/project/get.full"
                project_response = self._make_request(
                    project_detail_url,
                    {'access_token': self.goal_token, 'id': project['id']},
                    f"fetching project {project['name']}"
                )
                
                project_data = project_response.json()
                tasks = project_data.get('tasks', []) + project_data.get('subtasks', [])
                
                # Filter tasks by employee and time period
                employee_tasks_in_period = [
                    task for task in tasks 
                    if (task.get('username') == username and 
                        self._is_task_in_time_period(task, start_date, now))
                ]
                
                if employee_tasks_in_period:
                    employee_projects.append({
                        'project': project,
                        'tasks': employee_tasks_in_period,
                        'total_project_tasks': len(tasks),
                        'time_period': time_filter,
                        'project_data': project_data
                    })
            
            return employee_projects
            
        except Exception as e:
            st.error(f"Error fetching employee projects: {e}")
            return []

    def get_all_projects_data(self) -> List[Dict]:
        """Get all projects data for 360 analysis"""
        try:
            projects_url = "https://wework.base.vn/extapi/v3/project/list"
            projects_response = self._make_request(
                projects_url, 
                {'access_token': self.goal_token},
                "fetching all projects"
            )
            projects = projects_response.json().get('projects', [])
            
            all_projects_data = []
            
            for project in projects:
                project_detail_url = "https://wework.base.vn/extapi/v3/project/get.full"
                project_response = self._make_request(
                    project_detail_url,
                    {'access_token': self.goal_token, 'id': project['id']},
                    f"fetching project {project['name']}"
                )
                
                project_data = project_response.json()
                all_projects_data.append(project_data)
            
            return all_projects_data
            
        except Exception as e:
            st.error(f"Error fetching all projects: {e}")
            return []

    def _calculate_start_date(self, current_date: datetime, time_filter: str) -> datetime:
        """Calculate start date based on time filter"""
        if time_filter == "Tháng hiện tại":
            return current_date.replace(day=1)
        elif time_filter == "Quý hiện tại":
            quarter_start_month = ((current_date.month - 1) // 3) * 3 + 1
            return current_date.replace(month=quarter_start_month, day=1)
        elif time_filter == "Năm hiện tại":
            return current_date.replace(month=1, day=1)
        elif time_filter == "3 tháng gần nhất":
            return current_date - timedelta(days=90)
        elif time_filter == "6 tháng gần nhất":
            return current_date - timedelta(days=180)
        else:  # "Tất cả"
            return datetime(2020, 1, 1)

    def _is_task_in_time_period(self, task: dict, start_date: datetime, end_date: datetime) -> bool:
        """Check if task falls within time period"""
        try:
            # Check task creation time
            if task.get('start_time'):
                task_start = datetime.fromtimestamp(int(task['start_time']))
                if start_date <= task_start <= end_date:
                    return True
            
            # Check task completion time (for completed tasks)
            if task.get('completed_time') and int(task.get('completed_time', 0)) != 0:
                task_completion = datetime.fromtimestamp(int(task['completed_time']))
                if start_date <= task_completion <= end_date:
                    return True
            
            # Check task deadline (for all tasks with deadlines)
            if task.get('deadline') and int(task.get('deadline', 0)) != 0:
                task_deadline = datetime.fromtimestamp(int(task['deadline']))
                if start_date <= task_deadline <= end_date:
                    return True
            
            # Include tasks that are ongoing during the time period
            if task.get('start_time'):
                task_start = datetime.fromtimestamp(int(task['start_time']))
                if task_start < start_date:
                    completed_time = task.get('completed_time', 0)
                    if not completed_time or int(completed_time) == 0:
                        return True
                    else:
                        task_completion = datetime.fromtimestamp(int(completed_time))
                        if task_completion >= start_date:
                            return True
            
            # For time filter "Tất cả", include all tasks of the employee
            if start_date.year <= 2020:
                return True
                    
            return False
        except (ValueError, TypeError):
            return False

class Employee360Analyzer:
    """Advanced 360-degree employee analysis"""
    
    def __init__(self, api_client: APIClient, employees_df: pd.DataFrame):
        self.api_client = api_client
        self.employees_df = employees_df
        self.all_projects_data = []
        self.collaboration_network = {}
        self.skill_matrix = {}
        
    def load_organization_data(self):
        """Load complete organization data for 360 analysis"""
        with st.spinner("Đang tải dữ liệu tổ chức cho phân tích 360..."):
            self.all_projects_data = self.api_client.get_all_projects_data()
            self._build_collaboration_network()
            self._build_skill_matrix()
    
    def _build_collaboration_network(self):
        """Build collaboration network from all projects"""
        collaboration_counts = defaultdict(lambda: defaultdict(int))
        
        for project_data in self.all_projects_data:
            tasks = project_data.get('tasks', []) + project_data.get('subtasks', [])
            
            # Track collaboration between employees
            for task in tasks:
                username = task.get('username', '')
                if not username:
                    continue
                    
                followers = task.get('followers', [])
                for follower in followers:
                    follower_username = follower.get('username', '')
                    if follower_username and follower_username != username:
                        collaboration_counts[username][follower_username] += 1
        
        self.collaboration_network = dict(collaboration_counts)
    
    def _build_skill_matrix(self):
        """Build skill matrix based on task categories and project types"""
        employee_skills = defaultdict(lambda: defaultdict(int))
        
        for project_data in self.all_projects_data:
            project_name = project_data.get('name', '')
            tasks = project_data.get('tasks', []) + project_data.get('subtasks', [])
            
            for task in tasks:
                username = task.get('username', '')
                if not username:
                    continue
                
                # Extract skill categories
                task_category = ''
                if isinstance(task.get('tasklist'), dict):
                    task_category = task['tasklist'].get('name', '')
                
                if task_category:
                    employee_skills[username][task_category] += 1
                
                # Project-based skills
                if project_name:
                    employee_skills[username][f"Project: {project_name}"] += 1
        
        self.skill_matrix = dict(employee_skills)
    
    def analyze_employee_360(self, username: str, time_filter: str = "Tất cả") -> Dict:
        """Comprehensive 360-degree analysis of employee"""
        
        # Get employee's projects
        employee_projects = self.api_client.get_employee_projects_by_time(username, time_filter)
        
        if not employee_projects:
            return self._create_empty_360_analysis(username)
        
        # Core analysis components
        analysis = {
            'employee_username': username,
            'time_filter': time_filter,
            'collaboration_analysis': self._analyze_collaboration(username),
            'skill_analysis': self._analyze_skills(username),
            'performance_comparison': self._compare_with_peers(username, employee_projects),
            'leadership_influence': self._analyze_leadership_influence(username),
            'cross_functional_contribution': self._analyze_cross_functional_work(username),
            'growth_trajectory': self._analyze_growth_trajectory(username),
            'team_impact': self._analyze_team_impact(username),
            'innovation_contribution': self._analyze_innovation_metrics(username),
            'risk_assessment': self._assess_employee_risks(username, employee_projects),
            'recommendations_360': self._generate_360_recommendations(username, employee_projects)
        }
        
        return analysis
    
    def _analyze_collaboration(self, username: str) -> Dict:
        """Analyze collaboration patterns"""
        collaborations = self.collaboration_network.get(username, {})
        
        if not collaborations:
            return {
                'total_collaborators': 0,
                'frequent_collaborators': [],
                'collaboration_diversity': 0,
                'networking_score': 0
            }
        
        total_collaborators = len(collaborations)
        frequent_collaborators = [(user, count) for user, count in 
                                sorted(collaborations.items(), key=lambda x: x[1], reverse=True)[:5]]
        
        # Calculate diversity (how evenly distributed collaborations are)
        collaboration_values = list(collaborations.values())
        diversity_score = len(set(collaboration_values)) / len(collaboration_values) if collaboration_values else 0
        
        # Networking score based on total unique collaborators and interaction frequency
        networking_score = min(10, (total_collaborators * 0.5) + (sum(collaboration_values) * 0.1))
        
        return {
            'total_collaborators': total_collaborators,
            'frequent_collaborators': frequent_collaborators,
            'collaboration_diversity': round(diversity_score, 2),
            'networking_score': round(networking_score, 2),
            'total_interactions': sum(collaboration_values)
        }
    
    def _analyze_skills(self, username: str) -> Dict:
        """Analyze skill profile and expertise areas"""
        skills = self.skill_matrix.get(username, {})
        
        if not skills:
            return {
                'skill_areas': [],
                'expertise_level': 0,
                'skill_diversity': 0,
                'top_skills': []
            }
        
        # Calculate skill metrics
        total_skill_instances = sum(skills.values())
        skill_diversity = len(skills)
        
        # Top skills by frequency
        top_skills = [(skill, count) for skill, count in 
                     sorted(skills.items(), key=lambda x: x[1], reverse=True)[:10]]
        
        # Expertise level based on total experience
        expertise_level = min(10, total_skill_instances * 0.05)
        
        return {
            'skill_areas': list(skills.keys()),
            'expertise_level': round(expertise_level, 2),
            'skill_diversity': skill_diversity,
            'top_skills': top_skills,
            'total_skill_instances': total_skill_instances
        }
    
    def _compare_with_peers(self, username: str, employee_projects: List[Dict]) -> Dict:
        """Compare employee performance with organizational peers"""
        
        # Get employee's job title for peer comparison
        employee_info = self.employees_df[self.employees_df['username'] == username]
        if employee_info.empty:
            return {'peer_comparison': 'No peer data available'}
        
        job_title = employee_info.iloc[0]['job']
        
        # Find peers with similar job titles
        peers = self.employees_df[self.employees_df['job'] == job_title]['username'].tolist()
        peers = [p for p in peers if p != username]  # Exclude current employee
        
        if not peers:
            return {'peer_comparison': 'No peers found with similar job title'}
        
        # Calculate employee metrics
        employee_stats = self._calculate_employee_metrics(employee_projects)
        
        # Calculate peer metrics
        peer_metrics = []
        for peer in peers[:10]:  # Limit to top 10 peers for performance
            peer_projects = self.api_client.get_employee_projects_by_time(peer, "Tất cả")
            if peer_projects:
                peer_stats = self._calculate_employee_metrics(peer_projects)
                peer_metrics.append(peer_stats)
        
        if not peer_metrics:
            return {'peer_comparison': 'No peer performance data available'}
        
        # Calculate percentiles
        peer_completion_rates = [m['completion_rate'] for m in peer_metrics]
        peer_project_counts = [m['project_count'] for m in peer_metrics]
        
        completion_percentile = self._calculate_percentile(employee_stats['completion_rate'], peer_completion_rates)
        project_percentile = self._calculate_percentile(employee_stats['project_count'], peer_project_counts)
        
        return {
            'peer_count': len(peers),
            'completion_rate_percentile': completion_percentile,
            'project_count_percentile': project_percentile,
            'peer_average_completion': round(np.mean(peer_completion_rates), 2) if peer_completion_rates else 0,
            'employee_completion_rate': employee_stats['completion_rate'],
            'relative_performance': 'Above Average' if completion_percentile > 50 else 'Below Average'
        }
    
    def _calculate_employee_metrics(self, employee_projects: List[Dict]) -> Dict:
        """Calculate key metrics for an employee"""
        total_tasks = sum(len(proj['tasks']) for proj in employee_projects)
        completed_tasks = 0
        
        for project_info in employee_projects:
            for task in project_info['tasks']:
                if task.get('complete') == '100.00':
                    completed_tasks += 1
        
        completion_rate = Utils.calculate_percentage(completed_tasks, total_tasks)
        
        return {
            'project_count': len(employee_projects),
            'total_tasks': total_tasks,
            'completed_tasks': completed_tasks,
            'completion_rate': round(completion_rate, 2)
        }
    
    def _calculate_percentile(self, value: float, peer_values: List[float]) -> float:
        """Calculate percentile rank of value among peer values"""
        if not peer_values:
            return 50.0
        
        peer_values_sorted = sorted(peer_values)
        rank = sum(1 for v in peer_values_sorted if v <= value)
        percentile = (rank / len(peer_values_sorted)) * 100
        return round(percentile, 1)
    
    def _analyze_leadership_influence(self, username: str) -> Dict:
        """Analyze leadership and influence metrics"""
        
        # Count how many people this employee collaborates with (influence reach)
        collaborations = self.collaboration_network.get(username, {})
        influence_reach = len(collaborations)
        
        # Count how many times others seek this employee's collaboration
        reverse_collaborations = 0
        for other_user, their_collaborations in self.collaboration_network.items():
            if username in their_collaborations:
                reverse_collaborations += their_collaborations[username]
        
        # Leadership score based on influence metrics
        leadership_score = min(10, (influence_reach * 0.3) + (reverse_collaborations * 0.1))
        
        return {
            'influence_reach': influence_reach,
            'sought_after_score': reverse_collaborations,
            'leadership_score': round(leadership_score, 2),
            'leadership_level': self._get_leadership_level(leadership_score)
        }
    
    def _get_leadership_level(self, score: float) -> str:
        """Convert leadership score to descriptive level"""
        if score >= 8:
            return "Lãnh đạo xuất sắc"
        elif score >= 6:
            return "Có tiềm năng lãnh đạo"
        elif score >= 4:
            return "Đang phát triển kỹ năng lãnh đạo"
        else:
            return "Cần phát triển kỹ năng lãnh đạo"
    
    def _analyze_cross_functional_work(self, username: str) -> Dict:
        """Analyze cross-functional contribution"""
        skills = self.skill_matrix.get(username, {})
        
        # Count unique skill categories and project types
        skill_categories = [skill for skill in skills.keys() if not skill.startswith('Project:')]
        project_types = [skill for skill in skills.keys() if skill.startswith('Project:')]
        
        cross_functional_score = min(10, len(skill_categories) * 0.5 + len(project_types) * 0.3)
        
        return {
            'skill_categories': len(skill_categories),
            'project_types': len(project_types),
            'cross_functional_score': round(cross_functional_score, 2),
            'versatility_level': self._get_versatility_level(cross_functional_score)
        }
    
    def _get_versatility_level(self, score: float) -> str:
        """Convert versatility score to descriptive level"""
        if score >= 8:
            return "Rất linh hoạt đa chức năng"
        elif score >= 6:
            return "Khá linh hoạt"
        elif score >= 4:
            return "Có khả năng đa chức năng"
        else:
            return "Chuyên môn hóa cao"
    
    def _analyze_growth_trajectory(self, username: str) -> Dict:
        """Analyze employee growth over time"""
        
        # Get projects from different time periods
        current_quarter = self.api_client.get_employee_projects_by_time(username, "Quý hiện tại")
        last_quarter = self.api_client.get_employee_projects_by_time(username, "3 tháng gần nhất")
        
        current_metrics = self._calculate_employee_metrics(current_quarter) if current_quarter else {'completion_rate': 0, 'project_count': 0}
        previous_metrics = self._calculate_employee_metrics(last_quarter) if last_quarter else {'completion_rate': 0, 'project_count': 0}
        
        # Calculate growth rates
        completion_growth = current_metrics['completion_rate'] - previous_metrics['completion_rate']
        project_growth = current_metrics['project_count'] - previous_metrics['project_count']
        
        growth_trend = "Tăng trưởng" if completion_growth > 5 else "Ổn định" if completion_growth > -5 else "Cần cải thiện"
        
        return {
            'completion_rate_change': round(completion_growth, 2),
            'project_count_change': project_growth,
            'growth_trend': growth_trend,
            'trajectory_score': min(10, 5 + (completion_growth * 0.2))
        }
    
    def _analyze_team_impact(self, username: str) -> Dict:
        """Analyze impact on team and organization"""
        
        collaborations = self.collaboration_network.get(username, {})
        skills = self.skill_matrix.get(username, {})
        
        # Team impact metrics
        team_reach = len(collaborations)
        knowledge_sharing = sum(collaborations.values())  # Total interactions
        expertise_areas = len([skill for skill in skills.keys() if skills[skill] > 5])  # High-expertise areas
        
        impact_score = min(10, (team_reach * 0.3) + (knowledge_sharing * 0.05) + (expertise_areas * 0.5))
        
        return {
            'team_reach': team_reach,
            'knowledge_sharing_instances': knowledge_sharing,
            'expertise_areas': expertise_areas,
            'team_impact_score': round(impact_score, 2),
            'impact_level': self._get_impact_level(impact_score)
        }
    
    def _get_impact_level(self, score: float) -> str:
        """Convert impact score to descriptive level"""
        if score >= 8:
            return "Tác động cao đến tổ chức"
        elif score >= 6:
            return "Tác động tích cực"
        elif score >= 4:
            return "Tác động trung bình"
        else:
            return "Cần tăng cường tác động"
    
    def _analyze_innovation_metrics(self, username: str) -> Dict:
        """Analyze innovation and creativity metrics"""
        
        skills = self.skill_matrix.get(username, {})
        
        # Innovation indicators
        unique_projects = len([skill for skill in skills.keys() if skill.startswith('Project:')])
        diverse_skills = len(skills)
        
        # Innovation score based on project diversity and skill breadth
        innovation_score = min(10, (unique_projects * 0.4) + (diverse_skills * 0.2))
        
        return {
            'project_diversity': unique_projects,
            'skill_breadth': diverse_skills,
            'innovation_score': round(innovation_score, 2),
            'innovation_level': self._get_innovation_level(innovation_score)
        }
    
    def _get_innovation_level(self, score: float) -> str:
        """Convert innovation score to descriptive level"""
        if score >= 8:
            return "Rất sáng tạo và đổi mới"
        elif score >= 6:
            return "Có khả năng đổi mới"
        elif score >= 4:
            return "Tiềm năng sáng tạo"
        else:
            return "Cần phát triển tư duy đổi mới"
    
    def _assess_employee_risks(self, username: str, employee_projects: List[Dict]) -> Dict:
        """Assess potential risks related to employee"""
        
        risks = []
        risk_score = 0
        
        # Collaboration dependency risk
        collaborations = self.collaboration_network.get(username, {})
        if len(collaborations) < 3:
            risks.append("Hạn chế trong hợp tác nhóm")
            risk_score += 2
        
        # Skill concentration risk
        skills = self.skill_matrix.get(username, {})
        if len(skills) < 5:
            risks.append("Kỹ năng quá tập trung, thiếu đa dạng")
            risk_score += 2
        
        # Performance consistency risk
        if employee_projects:
            completion_rates = []
            for project_info in employee_projects:
                project_completed = sum(1 for task in project_info['tasks'] if task.get('complete') == '100.00')
                project_total = len(project_info['tasks'])
                if project_total > 0:
                    completion_rates.append(project_completed / project_total * 100)
            
            if completion_rates and np.std(completion_rates) > 30:
                risks.append("Hiệu suất không ổn định giữa các dự án")
                risk_score += 1
        
        # Engagement risk
        total_tasks = sum(len(proj['tasks']) for proj in employee_projects)
        if total_tasks < 10:
            risks.append("Mức độ tham gia thấp")
            risk_score += 2
        
        return {
            'identified_risks': risks,
            'risk_score': risk_score,
            'risk_level': self._get_risk_level(risk_score)
        }
    
    def _get_risk_level(self, score: int) -> str:
        """Convert risk score to descriptive level"""
        if score >= 6:
            return "Rủi ro cao"
        elif score >= 3:
            return "Rủi ro trung bình"
        else:
            return "Rủi ro thấp"
    
    def _generate_360_recommendations(self, username: str, employee_projects: List[Dict]) -> List[str]:
        """Generate comprehensive recommendations for employee development"""
        
        recommendations = []
        
        # Based on collaboration analysis
        collaborations = self.collaboration_network.get(username, {})
        if len(collaborations) < 5:
            recommendations.append("Tăng cường hợp tác và networking với đồng nghiệp")
        
        # Based on skill analysis
        skills = self.skill_matrix.get(username, {})
        if len(skills) < 8:
            recommendations.append("Mở rộng kỹ năng qua tham gia các dự án đa dạng")
        
        # Based on performance metrics
        if employee_projects:
            total_tasks = sum(len(proj['tasks']) for proj in employee_projects)
            completed_tasks = sum(1 for proj in employee_projects for task in proj['tasks'] if task.get('complete') == '100.00')
            completion_rate = Utils.calculate_percentage(completed_tasks, total_tasks)
            
            if completion_rate < 70:
                recommendations.append("Cải thiện hiệu suất hoàn thành công việc")
            elif completion_rate > 90:
                recommendations.append("Duy trì hiệu suất cao và chia sẻ kinh nghiệm")
        
        # Leadership development
        leadership_metrics = self._analyze_leadership_influence(username)
        if leadership_metrics['leadership_score'] < 5:
            recommendations.append("Phát triển kỹ năng lãnh đạo và ảnh hưởng")
        
        # Innovation and growth
        innovation_metrics = self._analyze_innovation_metrics(username)
        if innovation_metrics['innovation_score'] < 6:
            recommendations.append("Khuyến khích tư duy sáng tạo và đổi mới")
        
        return recommendations if recommendations else ["Tiếp tục duy trì hiệu suất tốt và phát triển toàn diện"]
    
    def _create_empty_360_analysis(self, username: str) -> Dict:
        """Create empty 360 analysis structure"""
        return {
            'employee_username': username,
            'time_filter': 'N/A',
            'collaboration_analysis': {'total_collaborators': 0, 'networking_score': 0},
            'skill_analysis': {'expertise_level': 0, 'skill_diversity': 0},
            'performance_comparison': {'peer_comparison': 'No data available'},
            'leadership_influence': {'leadership_score': 0, 'leadership_level': 'Không có dữ liệu'},
            'cross_functional_contribution': {'cross_functional_score': 0, 'versatility_level': 'Không có dữ liệu'},
            'growth_trajectory': {'growth_trend': 'Không có dữ liệu', 'trajectory_score': 0},
            'team_impact': {'team_impact_score': 0, 'impact_level': 'Không có dữ liệu'},
            'innovation_contribution': {'innovation_score': 0, 'innovation_level': 'Không có dữ liệu'},
            'risk_assessment': {'risk_level': 'Không đánh giá được', 'identified_risks': []},
            'recommendations_360': ['Cần thu thập thêm dữ liệu để phân tích']
        }

class ProjectAnalyzer:
    """Comprehensive project analysis for individual projects"""

    def __init__(self, api_client: APIClient):
        self.api_client = api_client

    def analyze_single_project(self, project_data: Dict, employee_username: str = None) -> Dict:
        """Analyze a single project in detail"""
        project = project_data.get('project', project_data)
        tasks = project_data.get('tasks', []) + project_data.get('subtasks', [])

        analysis = {
            'project_info': self._extract_project_info(project),
            'overall_stats': self._calculate_project_stats(tasks),
            'employee_contribution': self._analyze_employee_contribution(tasks, employee_username) if employee_username else None,
            'task_analysis': self._analyze_task_distribution(tasks),
            'timeline_analysis': self._analyze_project_timeline(tasks),
            'quality_metrics': self._analyze_quality_metrics(tasks),
            'collaboration_analysis': self._analyze_collaboration_in_project(tasks),
            'risk_assessment': self._assess_project_risks(tasks),
            'recommendations': self._generate_project_recommendations(tasks)
        }

        return analysis

    def _extract_project_info(self, project: Dict) -> Dict:
        """Extract key project information"""
        return {
            'id': project.get('id', ''),
            'name': project.get('name', 'Unknown Project'),
            'description': project.get('description', ''),
            'start_time': project.get('start_time', ''),
            'deadline': project.get('deadline', ''),
            'status': project.get('status', 'Unknown'),
            'priority': project.get('priority', 'Normal'),
            'progress': project.get('progress', '0')
        }

    def _calculate_project_stats(self, tasks: List[Dict]) -> Dict:
        """Calculate overall project statistics"""
        if not tasks:
            return self._get_empty_project_stats()

        total_tasks = len(tasks)
        completed_tasks = sum(1 for task in tasks if task.get('complete') == '100.00')
        ongoing_tasks = total_tasks - completed_tasks
        failed_tasks = sum(1 for task in tasks if isinstance(task.get('data'), dict) and task['data'].get('failed_reason'))

        # Calculate completion rate
        completion_rate = Utils.calculate_percentage(completed_tasks, total_tasks)

        # Calculate average completion percentage
        completion_percentages = [float(task.get('complete', '0')) for task in tasks]
        avg_completion = sum(completion_percentages) / len(completion_percentages) if completion_percentages else 0

        # Calculate overdue tasks
        overdue_tasks = 0
        on_time_tasks = 0

        for task in tasks:
            if task.get('complete') == '100.00' and task.get('completed_time') and task.get('deadline'):
                try:
                    deadline = datetime.fromtimestamp(int(task.get('deadline', 0)))
                    completed = datetime.fromtimestamp(int(task.get('completed_time')))
                    if completed <= deadline:
                        on_time_tasks += 1
                    else:
                        overdue_tasks += 1
                except:
                    pass

        return {
            'total_tasks': total_tasks,
            'completed_tasks': completed_tasks,
            'ongoing_tasks': ongoing_tasks,
            'failed_tasks': failed_tasks,
            'completion_rate': round(completion_rate, 2),
            'avg_completion_percentage': round(avg_completion, 2),
            'on_time_tasks': on_time_tasks,
            'overdue_tasks': overdue_tasks,
            'on_time_rate': round((on_time_tasks / completed_tasks * 100), 2) if completed_tasks > 0 else 0
        }

    def _get_empty_project_stats(self) -> Dict:
        """Return empty project statistics"""
        return {
            'total_tasks': 0,
            'completed_tasks': 0,
            'ongoing_tasks': 0,
            'failed_tasks': 0,
            'completion_rate': 0,
            'avg_completion_percentage': 0,
            'on_time_tasks': 0,
            'overdue_tasks': 0,
            'on_time_rate': 0
        }

    def _analyze_employee_contribution(self, tasks: List[Dict], employee_username: str) -> Dict:
        """Analyze specific employee's contribution to the project"""
        employee_tasks = [task for task in tasks if task.get('username') == employee_username]

        if not employee_tasks:
            return {'message': 'Employee not found in this project'}

        total_employee_tasks = len(employee_tasks)
        completed_employee_tasks = sum(1 for task in employee_tasks if task.get('complete') == '100.00')
        employee_completion_rate = (completed_employee_tasks / total_employee_tasks * 100) if total_employee_tasks > 0 else 0

        # Employee's task categories
        employee_categories = {}
        for task in employee_tasks:
            category = 'Chưa phân loại'
            if isinstance(task.get('tasklist'), dict):
                category = task['tasklist'].get('name', 'Chưa phân loại')
            employee_categories[category] = employee_categories.get(category, 0) + 1

        # Employee's contribution percentage to project
        contribution_percentage = (total_employee_tasks / len(tasks) * 100) if tasks else 0

        return {
            'total_tasks_assigned': total_employee_tasks,
            'completed_tasks': completed_employee_tasks,
            'completion_rate': round(employee_completion_rate, 2),
            'contribution_percentage': round(contribution_percentage, 2),
            'task_categories': employee_categories,
            'task_details': self._get_employee_task_details(employee_tasks)
        }

    def _get_employee_task_details(self, employee_tasks: List[Dict]) -> List[Dict]:
        """Get detailed information about employee's tasks"""
        task_details = []

        for task in employee_tasks:
            task_info = {
                'name': task.get('name', ''),
                'status': 'Hoàn thành' if task.get('complete') == '100.00' else 'Đang thực hiện',
                'category': task['tasklist'].get('name', 'Chưa phân loại') if isinstance(task.get('tasklist'), dict) else 'Chưa phân loại',
                'start_date': TaskAnalyzer.convert_timestamp(task.get('start_time')),
                'deadline': TaskAnalyzer.convert_timestamp(task.get('deadline')),
                'completion_date': TaskAnalyzer.convert_timestamp(task.get('completed_time')),
                'completion_percentage': task.get('complete', '0')
            }
            task_details.append(task_info)

        return task_details

    def _analyze_task_distribution(self, tasks: List[Dict]) -> Dict:
        """Analyze task distribution by categories and priorities"""
        categories = {}
        priorities = {}
        assignees = {}

        for task in tasks:
            # Category analysis
            category = 'Chưa phân loại'
            if isinstance(task.get('tasklist'), dict):
                category = task['tasklist'].get('name', 'Chưa phân loại')
            categories[category] = categories.get(category, 0) + 1

            # Priority analysis (if available)
            priority = task.get('priority', 'Normal')
            priorities[priority] = priorities.get(priority, 0) + 1

            # Assignee analysis
            assignee = task.get('username', 'Unassigned')
            assignees[assignee] = assignees.get(assignee, 0) + 1

        return {
            'by_category': categories,
            'by_priority': priorities,
            'by_assignee': assignees
        }

    def _analyze_project_timeline(self, tasks: List[Dict]) -> Dict:
        """Analyze project timeline and deadlines"""
        timeline_data = []

        for task in tasks:
            task_timeline = {
                'task_name': task.get('name', ''),
                'start_date': TaskAnalyzer.convert_timestamp(task.get('start_time')),
                'deadline': TaskAnalyzer.convert_timestamp(task.get('deadline')),
                'completion_date': TaskAnalyzer.convert_timestamp(task.get('completed_time')),
                'status': 'Hoàn thành' if task.get('complete') == '100.00' else 'Đang thực hiện',
                'is_overdue': False
            }

            # Check if overdue
            if (task.get('deadline') and
                task.get('complete') != '100.00' and
                int(task.get('deadline', 0)) != 0):
                try:
                    deadline = datetime.fromtimestamp(int(task.get('deadline')))
                    if datetime.now() > deadline:
                        task_timeline['is_overdue'] = True
                except:
                    pass

            timeline_data.append(task_timeline)

        return {
            'tasks': timeline_data,
            'upcoming_deadlines': self._get_upcoming_deadlines(timeline_data),
            'overdue_tasks': [t for t in timeline_data if t['is_overdue']]
        }

    def _get_upcoming_deadlines(self, timeline_data: List[Dict]) -> List[Dict]:
        """Get tasks with upcoming deadlines"""
        upcoming = []

        for task in timeline_data:
            if (task['deadline'] and
                task['status'] != 'Hoàn thành' and
                not task['is_overdue']):
                try:
                    deadline_date = datetime.strptime(task['deadline'], '%Y-%m-%d')
                    days_until_deadline = (deadline_date - datetime.now()).days

                    if 0 <= days_until_deadline <= 7:  # Within next 7 days
                        upcoming.append({
                            'task_name': task['task_name'],
                            'deadline': task['deadline'],
                            'days_until': days_until_deadline
                        })
                except:
                    pass

        return sorted(upcoming, key=lambda x: x['days_until'])

    def _analyze_quality_metrics(self, tasks: List[Dict]) -> Dict:
        """Analyze quality metrics of the project"""
        quality_metrics = {
            'tasks_with_description': 0,
            'tasks_with_deadline': 0,
            'tasks_with_result': 0,
            'tasks_with_followers': 0,
            'avg_description_length': 0
        }

        total_descriptions = 0
        description_lengths = []

        for task in tasks:
            # Check description quality
            content = task.get('content', '')
            if content and str(content).strip():
                quality_metrics['tasks_with_description'] += 1
                description_lengths.append(len(str(content)))

            # Check deadline
            if task.get('deadline') and int(task.get('deadline', 0)) != 0:
                quality_metrics['tasks_with_deadline'] += 1

            # Check result
            if isinstance(task.get('result'), dict) and task['result'].get('content', '').strip():
                quality_metrics['tasks_with_result'] += 1

            # Check followers
            if task.get('followers') and len(task['followers']) > 0:
                quality_metrics['tasks_with_followers'] += 1

        # Calculate average description length
        if description_lengths:
            quality_metrics['avg_description_length'] = round(sum(description_lengths) / len(description_lengths), 2)

        # Calculate percentages
        total_tasks = len(tasks)
        quality_metrics['description_rate'] = round(quality_metrics['tasks_with_description'] / total_tasks * 100, 2) if total_tasks > 0 else 0
        quality_metrics['deadline_rate'] = round(quality_metrics['tasks_with_deadline'] / total_tasks * 100, 2) if total_tasks > 0 else 0
        quality_metrics['result_rate'] = round(quality_metrics['tasks_with_result'] / total_tasks * 100, 2) if total_tasks > 0 else 0
        quality_metrics['follower_rate'] = round(quality_metrics['tasks_with_followers'] / total_tasks * 100, 2) if total_tasks > 0 else 0

        return quality_metrics

    def _analyze_collaboration_in_project(self, tasks: List[Dict]) -> Dict:
        """Analyze collaboration patterns within the project"""
        collaboration_matrix = {}

        for task in tasks:
            task_owner = task.get('username', 'Unknown')
            followers = task.get('followers', [])

            if task_owner not in collaboration_matrix:
                collaboration_matrix[task_owner] = {'followers': set(), 'following': set()}

            for follower in followers:
                follower_username = follower.get('username', '')
                if follower_username:
                    collaboration_matrix[task_owner]['followers'].add(follower_username)

                    if follower_username not in collaboration_matrix:
                        collaboration_matrix[follower_username] = {'followers': set(), 'following': set()}
                    collaboration_matrix[follower_username]['following'].add(task_owner)

        # Calculate collaboration metrics
        total_collaborations = sum(len(data['followers']) for data in collaboration_matrix.values())
        unique_collaborators = len([owner for owner, data in collaboration_matrix.items() if data['followers']])

        return {
            'collaboration_matrix': collaboration_matrix,
            'total_collaborations': total_collaborations,
            'unique_collaborators': unique_collaborators,
            'collaboration_density': round(total_collaborations / len(tasks), 2) if tasks else 0
        }

    def _assess_project_risks(self, tasks: List[Dict]) -> Dict:
        """Assess potential risks in the project"""
        risks = []
        risk_score = 0

        total_tasks = len(tasks)
        completed_tasks = sum(1 for task in tasks if task.get('complete') == '100.00')

        # Risk 1: Low completion rate
        if total_tasks > 0:
            completion_rate = completed_tasks / total_tasks * 100
            if completion_rate < 50:
                risks.append("Tỷ lệ hoàn thành thấp (< 50%)")
                risk_score += 3
            elif completion_rate < 75:
                risks.append("Tỷ lệ hoàn thành trung bình (50-75%)")
                risk_score += 1

        # Risk 2: Many overdue tasks
        overdue_tasks = sum(1 for task in tasks if self._is_task_overdue(task))
        if overdue_tasks > total_tasks * 0.3:
            risks.append(f"Quá nhiều task quá hạn ({overdue_tasks}/{total_tasks})")
            risk_score += 3
        elif overdue_tasks > total_tasks * 0.1:
            risks.append(f"Có task quá hạn ({overdue_tasks}/{total_tasks})")
            risk_score += 1

        # Risk 3: Poor data quality
        quality_tasks = sum(1 for task in tasks if task.get('content', '').strip())
        if quality_tasks / total_tasks < 0.7:
            risks.append("Nhiều task thiếu mô tả chi tiết")
            risk_score += 2

        # Risk 4: No collaboration
        collaboration_tasks = sum(1 for task in tasks if task.get('followers') and len(task['followers']) > 0)
        if collaboration_tasks / total_tasks < 0.3:
            risks.append("Thiếu sự hợp tác trong dự án")
            risk_score += 1

        return {
            'identified_risks': risks,
            'risk_score': risk_score,
            'risk_level': self._get_risk_level(risk_score)
        }

    def _is_task_overdue(self, task: Dict) -> bool:
        """Check if a task is overdue"""
        if (task.get('deadline') and
            task.get('complete') != '100.00' and
            int(task.get('deadline', 0)) != 0):
            try:
                deadline = datetime.fromtimestamp(int(task.get('deadline')))
                return datetime.now() > deadline
            except:
                pass
        return False

    def _get_risk_level(self, score: int) -> str:
        """Convert risk score to descriptive level"""
        return Utils.get_risk_level(score)

    def _generate_project_recommendations(self, tasks: List[Dict]) -> List[str]:
        """Generate recommendations for project improvement"""
        recommendations = []

        total_tasks = len(tasks)
        completed_tasks = sum(1 for task in tasks if task.get('complete') == '100.00')
        overdue_tasks = sum(1 for task in tasks if self._is_task_overdue(task))

        if completed_tasks / total_tasks < 0.7:
            recommendations.append("Tăng cường theo dõi và hỗ trợ hoàn thành task")

        if overdue_tasks > 0:
            recommendations.append("Ưu tiên xử lý các task quá hạn")

        # Check for tasks without descriptions
        tasks_without_desc = sum(1 for task in tasks if not task.get('content', '').strip())
        if tasks_without_desc / total_tasks > 0.3:
            recommendations.append("Cải thiện việc mô tả chi tiết cho các task")

        # Check collaboration
        collaboration_tasks = sum(1 for task in tasks if task.get('followers') and len(task['followers']) > 0)
        if collaboration_tasks / total_tasks < 0.5:
            recommendations.append("Khuyến khích hợp tác và giao tiếp trong team")

        if not recommendations:
            recommendations.append("Dự án đang tiến triển tốt, tiếp tục duy trì")

        return recommendations

class BasicAnalyzer:
    """Basic statistical analyzer without AI"""
    
    def __init__(self):
        self.required_fields = [
            'name', 'content', 'complete', 'start_time', 
            'username'
        ]
        
        self.important_missing_fields = [
            'start_time',
            'deadline',
            'tasklist'
        ]
        
        self.optional_fields = [
            'result', 'followers'
        ]
    
    def analyze_employee_performance_basic(self, employee_projects: List[Dict], employee_info: Dict, time_filter: str) -> Dict:
        """Basic statistical analysis of employee performance"""
        if not employee_projects:
            return self._create_empty_basic_analysis()
        
        analysis = {
            'employee_info': employee_info,
            'time_filter': time_filter,
            'projects_analysis': [],
            'overall_stats': {},
            'data_quality': {},
            'recommendations': []
        }
        
        # Analyze each project
        total_tasks = 0
        completed_tasks = 0
        failed_tasks = 0
        on_time_tasks = 0
        late_tasks = 0
        tasks_with_results = 0
        ongoing_tasks = 0
        tasks_with_deadlines = 0
        
        for project_info in employee_projects:
            project_analysis = self._analyze_project_data_quality(project_info)
            analysis['projects_analysis'].append(project_analysis)
            
            # Aggregate statistics
            tasks = project_info['tasks']
            total_tasks += len(tasks)
            
            for task in tasks:
                # Count tasks with deadlines
                if task.get('has_deadline') == '1' and task.get('deadline'):
                    tasks_with_deadlines += 1
                
                # Count completed tasks
                if task.get('complete') == '100.00':
                    completed_tasks += 1
                    
                    # Check deadline compliance
                    if task.get('has_deadline') == '1' and task.get('completed_time') and task.get('deadline'):
                        try:
                            deadline = datetime.fromtimestamp(int(task.get('deadline', 0)))
                            completed = datetime.fromtimestamp(int(task.get('completed_time')))
                            if completed <= deadline:
                                on_time_tasks += 1
                            else:
                                late_tasks += 1
                        except:
                            pass
                else:
                    ongoing_tasks += 1
                
                # Count tasks with results
                if isinstance(task.get('result'), dict) and task['result'].get('content', '').strip():
                    tasks_with_results += 1
                
                # Count failed tasks
                if isinstance(task.get('data'), dict):
                    failed_reason = task['data'].get('failed_reason', '')
                    if failed_reason:
                        failed_tasks += 1
        
        # Calculate overall statistics
        completion_rate = Utils.calculate_percentage(completed_tasks, total_tasks)
        on_time_rate = Utils.calculate_percentage(on_time_tasks, completed_tasks)
        on_time_completion_rate = Utils.calculate_percentage(on_time_tasks, total_tasks)
        result_documentation_rate = Utils.calculate_percentage(tasks_with_results, total_tasks)
        
        analysis['overall_stats'] = {
            'total_projects': len(employee_projects),
            'total_tasks': total_tasks,
            'completed_tasks': completed_tasks,
            'ongoing_tasks': ongoing_tasks,
            'failed_tasks': failed_tasks,
            'completion_rate': round(completion_rate, 2),
            'on_time_tasks': on_time_tasks,
            'late_tasks': late_tasks,
            'on_time_rate': round(on_time_rate, 2),
            'on_time_completion_rate': round(on_time_completion_rate, 2),
            'tasks_with_results': tasks_with_results,
            'result_documentation_rate': round(result_documentation_rate, 2),
            'tasks_with_deadlines': tasks_with_deadlines
        }
        
        # Generate basic recommendations
        analysis['recommendations'] = self._generate_basic_recommendations(analysis['overall_stats'])
        
        # Overall data quality assessment
        analysis['data_quality'] = self._assess_overall_data_quality(analysis['projects_analysis'])
        
        return analysis
    
    def _analyze_project_data_quality(self, project_info: Dict) -> Dict:
        """Analyze data quality for a single project"""
        project = project_info['project']
        tasks = project_info['tasks']
        
        quality_stats = {
            'project_name': project['name'],
            'project_id': project['id'],
            'total_tasks': len(tasks),
            'required_fields_complete': 0,
            'optional_fields_complete': 0,
            'missing_fields': defaultdict(int),
            'important_missing_fields': defaultdict(int),
            'quality_score': 0,
            'issues': []
        }
        
        for task in tasks:
            # Check required fields
            required_complete = 0
            for field in self.required_fields:
                if field in task and task[field] and str(task[field]).strip():
                    required_complete += 1
                else:
                    quality_stats['missing_fields'][field] += 1
            
            quality_stats['required_fields_complete'] += required_complete
            
            # Check important missing fields
            for field in self.important_missing_fields:
                if field == 'start_time':
                    if not task.get('start_time') or not str(task.get('start_time')).strip():
                        quality_stats['important_missing_fields']['Ngày bắt đầu'] += 1
                elif field == 'deadline':
                    if not task.get('deadline') or not str(task.get('deadline')).strip() or task.get('deadline') == '0':
                        quality_stats['important_missing_fields']['Deadline'] += 1
                elif field == 'tasklist':
                    if not isinstance(task.get('tasklist'), dict) or not task['tasklist'].get('name', '').strip():
                        quality_stats['important_missing_fields']['Loại công việc'] += 1
            
            # Check optional fields
            optional_complete = 0
            for field in self.optional_fields:
                if field in task and task[field]:
                    if field == 'result' and isinstance(task[field], dict):
                        if task[field].get('content', '').strip():
                            optional_complete += 1
                    elif field == 'followers' and isinstance(task[field], list):
                        if len(task[field]) > 0:
                            optional_complete += 1
                    else:
                        optional_complete += 1
            
            quality_stats['optional_fields_complete'] += optional_complete
        
        # Calculate quality score (0-100)
        if len(tasks) > 0:
            max_required = len(self.required_fields) * len(tasks)
            max_optional = len(self.optional_fields) * len(tasks)
            
            required_score = (quality_stats['required_fields_complete'] / max_required * 70) if max_required > 0 else 0
            optional_score = (quality_stats['optional_fields_complete'] / max_optional * 30) if max_optional > 0 else 0
            
            quality_stats['quality_score'] = round(required_score + optional_score, 2)
        
        # Identify common issues based on important missing fields
        if quality_stats['important_missing_fields']['Ngày bắt đầu'] > len(tasks) * 0.3:
            quality_stats['issues'].append("Nhiều task thiếu ngày bắt đầu")
        
        if quality_stats['important_missing_fields']['Deadline'] > len(tasks) * 0.5:
            quality_stats['issues'].append("Nhiều task không có deadline")
        
        if quality_stats['important_missing_fields']['Loại công việc'] > len(tasks) * 0.3:
            quality_stats['issues'].append("Nhiều task chưa được phân loại")
        
        if quality_stats['missing_fields']['content'] > len(tasks) * 0.3:
            quality_stats['issues'].append("Nhiều task thiếu mô tả chi tiết")
        
        if quality_stats['optional_fields_complete'] < len(tasks) * 0.2:
            quality_stats['issues'].append("Thiếu thông tin kết quả và người liên quan")
        
        return quality_stats
    
    def _assess_overall_data_quality(self, projects_analysis: List[Dict]) -> Dict:
        """Assess overall data quality across all projects"""
        if not projects_analysis:
            return {'score': 0, 'issues': ["Không có dữ liệu để đánh giá"]}
        
        total_score = sum(p['quality_score'] for p in projects_analysis)
        avg_score = total_score / len(projects_analysis)
        
        all_issues = []
        for project in projects_analysis:
            all_issues.extend(project['issues'])
        
        # Count most common issues
        issue_counts = defaultdict(int)
        for issue in all_issues:
            issue_counts[issue] += 1
        
        common_issues = [f"{issue} ({count} dự án)" for issue, count in issue_counts.items() if count > 1]
        
        return {
            'average_score': round(avg_score, 2),
            'grade': self._get_quality_grade(avg_score),
            'common_issues': common_issues,
            'projects_with_issues': len([p for p in projects_analysis if p['issues']]),
            'total_projects': len(projects_analysis)
        }
    
    def _get_quality_grade(self, score: float) -> str:
        """Convert quality score to grade"""
        return Utils.get_quality_grade(score)
    
    def _generate_basic_recommendations(self, stats: Dict) -> List[str]:
        """Generate basic recommendations based on statistics"""
        recommendations = []
        
        if stats['completion_rate'] < 70:
            recommendations.append("Cần cải thiện tỷ lệ hoàn thành công việc")
        elif stats['completion_rate'] > 90:
            recommendations.append("Duy trì tỷ lệ hoàn thành công việc tốt")
        
        if stats['on_time_completion_rate'] < 60:
            recommendations.append("Cần tăng cường quản lý thời gian và tuân thủ deadline")
        elif stats['on_time_completion_rate'] > 85:
            recommendations.append("Khả năng quản lý thời gian tốt")
        
        if stats['result_documentation_rate'] < 50:
            recommendations.append("Cần cải thiện việc ghi chép kết quả công việc")
        
        if stats['total_projects'] > 5:
            recommendations.append("Khả năng làm việc đa dự án tốt")
        
        if stats['failed_tasks'] > stats['total_tasks'] * 0.1:
            recommendations.append("Cần phân tích nguyên nhân thất bại và có biện pháp cải thiện")
        
        if stats['ongoing_tasks'] > stats['total_tasks'] * 0.3:
            recommendations.append("Có nhiều công việc đang thực hiện, cần ưu tiên hoàn thành")
        
        return recommendations if recommendations else ["Hiệu suất làm việc ổn định"]
    
    def _create_empty_basic_analysis(self) -> Dict:
        """Create empty analysis structure"""
        return {
            'employee_info': {},
            'time_filter': '',
            'projects_analysis': [],
            'overall_stats': {
                'total_projects': 0,
                'total_tasks': 0,
                'completed_tasks': 0,
                'ongoing_tasks': 0,
                'failed_tasks': 0,
                'completion_rate': 0,
                'on_time_tasks': 0,
                'late_tasks': 0,
                'on_time_rate': 0,
                'on_time_completion_rate': 0,
                'tasks_with_results': 0,
                'result_documentation_rate': 0,
                'tasks_with_deadlines': 0
            },
            'data_quality': {
                'average_score': 0,
                'grade': "Không có dữ liệu",
                'common_issues': [],
                'projects_with_issues': 0,
                'total_projects': 0
            },
            'recommendations': ["Không có dữ liệu để phân tích"]
        }

class AIAnalyzer:
    def __init__(self):
        self.api_key = GEMINI_API_KEY
        self.client = None
        self.use_gemma = False
        
        if GENAI_AVAILABLE and self.api_key:
            try:
                self.client = genai.Client(api_key=self.api_key)
                self._test_connection()
            except Exception as e:
                st.error(f"Lỗi khởi tạo AI: {e}")
                self.client = None

    def _test_connection(self):
        """Test AI connection and determine which model to use"""
        try:
            contents = [types.Content(role="user", parts=[types.Part.from_text(text="test")])]
            config = types.GenerateContentConfig(temperature=0.1)
            
            response_text = ""
            for chunk in self.client.models.generate_content_stream(
                model="gemini-2.0-flash",
                contents=contents,
                config=config,
            ):
                response_text += chunk.text
                break
            
            self.use_gemma = False
            st.success("✅ Gemini AI kết nối thành công")
            
        except Exception as e:
            st.warning(f"⚠️ Gemini không khả dụng ({str(e)}), chuyển sang Gemma...")
            try:
                contents = [types.Content(role="user", parts=[types.Part.from_text(text="test")])]
                config = types.GenerateContentConfig(temperature=0.1)
                
                response_text = ""
                for chunk in self.client.models.generate_content_stream(
                    model="gemma-2-27b-it",
                    contents=contents,
                    config=config,
                ):
                    response_text += chunk.text
                    break
                
                self.use_gemma = True
                st.success("✅ Gemma AI kết nối thành công")
                
            except Exception as e2:
                st.error(f"❌ Không thể kết nối AI: {str(e2)}")
                self.client = None

    def analyze_employee_performance_by_time(self, employee_projects: List[Dict], employee_info: Dict, time_filter: str) -> Dict:
        """Analyze employee performance across multiple projects in time period"""
        if not self.client:
            return None
        
        try:
            employee_summary = self._prepare_employee_time_data(employee_projects, employee_info, time_filter)
            
            prompt = f"""
            Phân tích hiệu suất nhân viên trong khoảng thời gian {time_filter} dựa trên dữ liệu sau:
            
            Thông tin nhân viên:
            - Tên: {employee_info.get('name', 'N/A')}
            - Username: {employee_info.get('username', 'N/A')}
            - Vị trí: {employee_info.get('job', 'N/A')}
            - Email: {employee_info.get('email', 'N/A')}
            
            {employee_summary}
            
            Hãy đánh giá nhân viên theo các tiêu chí sau và trả về kết quả {"JSON" if not self.use_gemma else "có cấu trúc"}:
            1. Hiệu suất làm việc tổng thể trong kỳ (1-10)
            2. Khả năng hoàn thành công việc (1-10)
            3. Chất lượng công việc (1-10)
            4. Tuân thủ deadline (1-10)
            5. Tính nhất quán trong công việc (1-10)
            6. Khả năng đa dự án (1-10)
            7. Xu hướng cải thiện trong kỳ (1-10)
            8. Mức độ đóng góp tổng thể (1-10)
            9. Điểm mạnh nổi bật (danh sách)
            10. Điểm cần cải thiện (danh sách)
            11. Khuyến nghị phát triển (danh sách)
            12. Đánh giá xu hướng hiệu suất
            13. So sánh với kỳ trước (nếu có thể đánh giá)
            14. Khuyến nghị cho kỳ tiếp theo
            
            {"Định dạng JSON với keys: hieu_suat_tong_the, kha_nang_hoan_thanh, chat_luong_cong_viec, tuan_thu_deadline, tinh_nhat_quan, kha_nang_da_du_an, xu_huong_cai_thien, dong_gop_tong_the, diem_manh, diem_can_cai_thien, khuyen_nghi_phat_trien, xu_huong_hieu_suat, so_sanh_ky_truoc, khuyen_nghi_ky_tiep" if not self.use_gemma else "Trả lời có cấu trúc rõ ràng với từng tiêu chí"}
            """
            
            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part.from_text(text=prompt)],
                ),
            ]
            
            if self.use_gemma:
                config = types.GenerateContentConfig(
                    temperature=0.7,
                )
                
                response_text = ""
                for chunk in self.client.models.generate_content_stream(
                    model="gemma-2-27b-it",
                    contents=contents,
                    config=config,
                ):
                    response_text += chunk.text
                
                return self._extract_json_from_employee_analysis(response_text)
                
            else:
                config = types.GenerateContentConfig(
                    temperature=0.7,
                    response_mime_type="application/json",
                    response_schema=genai.types.Schema(
                        type=genai.types.Type.OBJECT,
                        required=["hieu_suat_tong_the", "kha_nang_hoan_thanh", "chat_luong_cong_viec"],
                        properties={
                            "hieu_suat_tong_the": genai.types.Schema(type=genai.types.Type.INTEGER),
                            "kha_nang_hoan_thanh": genai.types.Schema(type=genai.types.Type.INTEGER),
                            "chat_luong_cong_viec": genai.types.Schema(type=genai.types.Type.INTEGER),
                            "tuan_thu_deadline": genai.types.Schema(type=genai.types.Type.INTEGER),
                            "tinh_nhat_quan": genai.types.Schema(type=genai.types.Type.INTEGER),
                            "kha_nang_da_du_an": genai.types.Schema(type=genai.types.Type.INTEGER),
                            "xu_huong_cai_thien": genai.types.Schema(type=genai.types.Type.INTEGER),
                            "dong_gop_tong_the": genai.types.Schema(type=genai.types.Type.INTEGER),
                            "diem_manh": genai.types.Schema(
                                type=genai.types.Type.ARRAY,
                                items=genai.types.Schema(type=genai.types.Type.STRING)
                            ),
                            "diem_can_cai_thien": genai.types.Schema(
                                type=genai.types.Type.ARRAY,
                                items=genai.types.Schema(type=genai.types.Type.STRING)
                            ),
                            "khuyen_nghi_phat_trien": genai.types.Schema(
                                type=genai.types.Type.ARRAY,
                                items=genai.types.Schema(type=genai.types.Type.STRING)
                            ),
                            "xu_huong_hieu_suat": genai.types.Schema(type=genai.types.Type.STRING),
                            "so_sanh_ky_truoc": genai.types.Schema(type=genai.types.Type.STRING),
                            "khuyen_nghi_ky_tiep": genai.types.Schema(type=genai.types.Type.STRING),
                        },
                    ),
                )
                
                response_text = ""
                for chunk in self.client.models.generate_content_stream(
                    model="gemini-2.0-flash",
                    contents=contents,
                    config=config,
                ):
                    response_text += chunk.text
                
                return json.loads(response_text)
                
        except Exception as e:
            st.error(f"Lỗi phân tích AI nhân viên: {e}")
            return self._create_default_employee_analysis()

    def _prepare_employee_time_data(self, employee_projects: List[Dict], employee_info: Dict, time_filter: str) -> str:
        """Prepare employee performance data for AI analysis"""
        if not employee_projects:
            return f"Không có dữ liệu công việc trong khoảng thời gian {time_filter}"
        
        total_tasks = 0
        completed_tasks = 0
        failed_tasks = 0
        ongoing_tasks = 0
        on_time_tasks = 0
        projects_count = len(employee_projects)
        
        project_summaries = []
        
        for project_info in employee_projects:
            project_name = project_info['project']['name']
            tasks = project_info['tasks']
            total_project_tasks = project_info['total_project_tasks']
            
            total_tasks += len(tasks)
            
            project_completed = 0
            project_failed = 0
            project_ongoing = 0
            project_on_time = 0
            
            for task in tasks:
                if task.get('complete') == '100.00':
                    completed_tasks += 1
                    project_completed += 1
                    
                    # Check if completed on time
                    if task.get('has_deadline') == '1' and task.get('completed_time'):
                        try:
                            deadline = datetime.fromtimestamp(int(task.get('deadline', 0)))
                            completed = datetime.fromtimestamp(int(task.get('completed_time')))
                            if completed <= deadline:
                                on_time_tasks += 1
                                project_on_time += 1
                        except:
                            pass
                else:
                    ongoing_tasks += 1
                    project_ongoing += 1
                
                # Check for failed tasks
                if isinstance(task.get('data'), dict):
                    failed_reason = task['data'].get('failed_reason', '')
                    if failed_reason:
                        failed_tasks += 1
                        project_failed += 1
            
            project_summaries.append(f"""
            - Dự án: {project_name}
              + Tham gia: {len(tasks)} công việc
              + Hoàn thành: {project_completed} ({project_completed/len(tasks)*100:.1f}%)
              + Đang thực hiện: {project_ongoing} ({project_ongoing/len(tasks)*100:.1f}%)
              + Thất bại: {project_failed}
              + Đúng hạn: {project_on_time}
            """)
        
        completion_rate = completed_tasks / total_tasks * 100 if total_tasks > 0 else 0
        ongoing_rate = ongoing_tasks / total_tasks * 100 if total_tasks > 0 else 0
        on_time_rate = on_time_tasks / completed_tasks * 100 if completed_tasks > 0 else 0
        
        return f"""
        Thống kê {time_filter}:
        - Tham gia: {projects_count} dự án
        - Tổng công việc: {total_tasks}
        - Hoàn thành: {completed_tasks} ({completion_rate:.1f}%)
        - Đang thực hiện: {ongoing_tasks} ({ongoing_rate:.1f}%)
        - Thất bại: {failed_tasks}
        - Hoàn thành đúng hạn: {on_time_tasks}/{completed_tasks} ({on_time_rate:.1f}%)
        
        Chi tiết theo dự án:
        {''.join(project_summaries)}
        """

    def _extract_json_from_employee_analysis(self, text):
        """Extract JSON from employee analysis text"""
        try:
            json_match = re.search(r'\{.*\}', text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            
            return self._create_default_employee_analysis()
        except:
            return self._create_default_employee_analysis()

    def _create_default_employee_analysis(self):
        """Create default employee analysis structure"""
        return {
            'hieu_suat_tong_the': 5,
            'kha_nang_hoan_thanh': 5,
            'chat_luong_cong_viec': 5,
            'tuan_thu_deadline': 5,
            'tinh_nhat_quan': 5,
            'kha_nang_da_du_an': 5,
            'xu_huong_cai_thien': 5,
            'dong_gop_tong_the': 5,
            'diem_manh': ["Có kinh nghiệm làm việc đa dự án", "Tham gia tích cực"],
            'diem_can_cai_thien': ["Cần cải thiện hiệu suất", "Cần tăng cường kỹ năng quản lý thời gian"],
            'khuyen_nghi_phat_trien': ["Tham gia khóa đào tạo", "Cải thiện quy trình làm việc"],
            'xu_huong_hieu_suat': "Ổn định với tiềm năng cải thiện",
            'so_sanh_ky_truoc': "Cần dữ liệu kỳ trước để so sánh chính xác",
            'khuyen_nghi_ky_tiep': "Tiếp tục duy trì hiệu suất và tập trung cải thiện các điểm yếu"
        }

class TaskAnalyzer:
    def __init__(self):
        pass
    
    @staticmethod
    def convert_timestamp(timestamp):
        """Convert timestamp to datetime string"""
        return Utils.safe_timestamp_convert(timestamp)

    @staticmethod
    def clean_html_content(content):
        """Clean HTML content and extract plain text"""
        return Utils.clean_html_content(content)

    def parse_employee_tasks(self, employee_projects: List[Dict]) -> pd.DataFrame:
        """Parse all tasks from employee's projects into a single DataFrame"""
        all_tasks_data = []
        
        for project_info in employee_projects:
            project_name = project_info['project']['name']
            tasks = project_info['tasks']
            
            for task in tasks:
                try:
                    content = self.clean_html_content(task.get('content', ''))
                    
                    task_category = ''
                    if isinstance(task.get('tasklist'), dict):
                        task_category = task['tasklist'].get('name', '')
                    
                    if not task_category or task_category.strip() == '':
                        task_category = 'Chưa phân loại'
                    
                    result = ''
                    if isinstance(task.get('result'), dict):
                        result = task['result'].get('content', '')
                    
                    failed_reason = ''
                    if isinstance(task.get('data'), dict):
                        data_dict = task['data']
                        if isinstance(data_dict.get('failed_reason'), dict):
                            failed_reason = data_dict['failed_reason'].get('reason', '')
                        elif isinstance(data_dict.get('failed_reason'), str):
                            failed_reason = data_dict['failed_reason']
        
                    deadline = ''
                    if int(task.get('has_deadline', '0')) == 1:
                        deadline = self.convert_timestamp(task.get('deadline'))
                    
                    completion_date = ''
                    if task.get('completed_time') and int(task.get('completed_time', 0)) != 0:
                        completion_date = self.convert_timestamp(task.get('completed_time'))
                    
                    if failed_reason:
                        status = 'Thất bại'
                    elif task.get('complete') == '100.00':
                        status = 'Hoàn thành'
                    else:
                        status = 'Đang thực hiện'
                    
                    created_date = self.convert_timestamp(task.get('start_time'))
                    
                    task_data = {
                        'Dự án': project_name,
                        'Loại công việc': task_category,
                        'Tên công việc': task.get('name', ''),
                        'Mô tả công việc': content,
                        'Trạng thái': status,
                        'Kết quả đạt được': result,
                        'Lí do thất bại': failed_reason,
                        'Ngày bắt đầu': created_date,
                        'Deadline': deadline,
                        'Ngày hoàn thành': completion_date,
                        'Người liên quan': ', '.join([follower.get('username', '') for follower in task.get('followers', [])]),
                        'Mức độ hoàn thành': f"{task.get('complete', '0')}%",
                    }
                    
                    all_tasks_data.append(task_data)
                    
                except Exception as e:
                    st.error(f"Error parsing task: {str(e)}")
                    continue
        
        return pd.DataFrame(all_tasks_data)

    def export_employee_report_to_word(self, df, employee_name, time_period, analysis_result=None, analysis_type="Cơ bản"):
        """Export employee report to Word document"""
        doc = Document()
        
        title = doc.add_heading(f'Báo cáo Hiệu suất Nhân viên: {employee_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        period_para = doc.add_paragraph(f'Khoảng thời gian: {time_period} | Phương pháp: {analysis_type}')
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        if analysis_result:
            if analysis_type == "AI" or "ai_analysis" in analysis_result:
                self._add_ai_analysis_to_word(doc, analysis_result)
            elif analysis_type == "Cơ bản" or "overall_stats" in analysis_result:
                self._add_basic_analysis_to_word(doc, analysis_result)
        
        projects = df['Dự án'].unique()
        
        for idx, project in enumerate(projects, start=1):
            project_tasks = df[df['Dự án'] == project]
            
            doc.add_heading(f"{idx}. {project}", level=1)
            
            total_tasks = len(project_tasks)
            completed_tasks = len(project_tasks[project_tasks['Trạng thái'] == 'Hoàn thành'])
            ongoing_tasks = len(project_tasks[project_tasks['Trạng thái'] == 'Đang thực hiện'])
            failed_tasks = len(project_tasks[project_tasks['Trạng thái'] == 'Thất bại'])
            completion_rate = completed_tasks / total_tasks * 100 if total_tasks > 0 else 0
            
            stats_para = doc.add_paragraph()
            stats_para.add_run('Thống kê: ').bold = True
            stats_para.add_run(f'Tổng số công việc: {total_tasks}, Hoàn thành: {completed_tasks} ({completion_rate:.1f}%), Đang thực hiện: {ongoing_tasks}, Thất bại: {failed_tasks}')
            
            for _, task in project_tasks.iterrows():
                doc.add_heading(task['Tên công việc'], level=2)
                
                if task['Mô tả công việc']:
                    p = doc.add_paragraph()
                    p.add_run('Mô tả: ').bold = True
                    p.add_run(task['Mô tả công việc'])
                
                p = doc.add_paragraph()
                p.add_run('Trạng thái: ').bold = True
                p.add_run(f"{task['Trạng thái']} ({task['Mức độ hoàn thành']})")
                
                if task['Ngày bắt đầu'] or task['Deadline'] or task['Ngày hoàn thành']:
                    p = doc.add_paragraph()
                    if task['Ngày bắt đầu']:
                        p.add_run('Ngày bắt đầu: ').bold = True
                        p.add_run(task['Ngày bắt đầu'])
                        if task['Deadline'] or task['Ngày hoàn thành']:
                            p.add_run(' | ')
                    
                    if task['Deadline']:
                        p.add_run('Deadline: ').bold = True
                        p.add_run(task['Deadline'])
                        if task['Ngày hoàn thành']:
                            p.add_run(' | ')
                    
                    if task['Ngày hoàn thành']:
                        p.add_run('Ngày hoàn thành: ').bold = True
                        p.add_run(task['Ngày hoàn thành'])
                
                if task['Kết quả đạt được'] and str(task['Kết quả đạt được']).strip():
                    p = doc.add_paragraph()
                    p.add_run('Kết quả: ').bold = True
                    p.add_run(task['Kết quả đạt được'])
                
                if task['Lí do thất bại'] and str(task['Lí do thất bại']).strip():
                    p = doc.add_paragraph()
                    p.add_run('Lý do thất bại: ').bold = True
                    p.add_run(task['Lí do thất bại'])
                
                doc.add_paragraph('─' * 50)
            
            doc.add_paragraph()
        
        return doc
    
    def _add_ai_analysis_to_word(self, doc, analysis_result):
        """Add AI analysis to Word document"""
        ai_result = analysis_result.get('ai_analysis', analysis_result)
        
        doc.add_heading('🤖 Đánh giá AI', level=1)
        
        doc.add_heading('Điểm số Hiệu suất:', level=2)
        metrics = [
            ('Hiệu suất tổng thể', ai_result.get('hieu_suat_tong_the', 0)),
            ('Khả năng hoàn thành', ai_result.get('kha_nang_hoan_thanh', 0)),
            ('Chất lượng công việc', ai_result.get('chat_luong_cong_viec', 0)),
            ('Tuân thủ deadline', ai_result.get('tuan_thu_deadline', 0)),
            ('Tính nhất quán', ai_result.get('tinh_nhat_quan', 0)),
            ('Khả năng đa dự án', ai_result.get('kha_nang_da_du_an', 0))
        ]
        
        for metric, score in metrics:
            p = doc.add_paragraph(f'• {metric}: {score}/10')
        
        doc.add_heading('🎯 Điểm mạnh:', level=2)
        for strength in ai_result.get('diem_manh', []):
            doc.add_paragraph(f'• {strength}')
        
        doc.add_heading('⚠️ Điểm cần cải thiện:', level=2)
        for improvement in ai_result.get('diem_can_cai_thien', []):
            doc.add_paragraph(f'• {improvement}')
        
        doc.add_heading('💡 Khuyến nghị:', level=2)
        for rec in ai_result.get('khuyen_nghi_phat_trien', []):
            doc.add_paragraph(f'• {rec}')
        
        doc.add_paragraph()
    
    def _add_basic_analysis_to_word(self, doc, analysis_result):
        """Add basic analysis to Word document"""
        stats = analysis_result.get('overall_stats', {})
        data_quality = analysis_result.get('data_quality', {})
        
        doc.add_heading('📊 Phân tích Cơ bản', level=1)
        
        doc.add_heading('Thống kê Tổng quan:', level=2)
        doc.add_paragraph(f'• Tổng số dự án: {stats.get("total_projects", 0)}')
        doc.add_paragraph(f'• Tổng công việc: {stats.get("total_tasks", 0)}')
        doc.add_paragraph(f'• Hoàn thành: {stats.get("completed_tasks", 0)}')
        doc.add_paragraph(f'• Đang thực hiện: {stats.get("ongoing_tasks", 0)}')
        doc.add_paragraph(f'• Tỷ lệ hoàn thành: {stats.get("completion_rate", 0)}%')
        doc.add_paragraph(f'• Tỷ lệ hoàn thành đúng hạn: {stats.get("on_time_completion_rate", 0)}%')
        doc.add_paragraph(f'• Tỷ lệ ghi chép kết quả: {stats.get("result_documentation_rate", 0)}%')
        
        doc.add_heading('Chất lượng Dữ liệu:', level=2)
        doc.add_paragraph(f'• Điểm chất lượng: {data_quality.get("average_score", 0)}/100')
        doc.add_paragraph(f'• Xếp hạng: {data_quality.get("grade", "N/A")}')
        
        doc.add_heading('Khuyến nghị:', level=2)
        for rec in analysis_result.get('recommendations', []):
            doc.add_paragraph(f'• {rec}')
        
        doc.add_paragraph()

# Visualization functions
def create_employee_performance_radar(analysis_result):
    """Create radar chart for employee performance analysis"""
    categories = [
        'Hiệu suất tổng thể',
        'Khả năng hoàn thành',
        'Chất lượng công việc',
        'Tuân thủ deadline',
        'Tính nhất quán',
        'Khả năng đa dự án',
        'Xu hướng cải thiện',
        'Đóng góp tổng thể'
    ]
    
    values = [
        analysis_result.get('hieu_suat_tong_the', 0),
        analysis_result.get('kha_nang_hoan_thanh', 0),
        analysis_result.get('chat_luong_cong_viec', 0),
        analysis_result.get('tuan_thu_deadline', 0),
        analysis_result.get('tinh_nhat_quan', 0),
        analysis_result.get('kha_nang_da_du_an', 0),
        analysis_result.get('xu_huong_cai_thien', 0),
        analysis_result.get('dong_gop_tong_the', 0)
    ]
    
    fig = go.Figure()
    
    fig.add_trace(go.Scatterpolar(
        r=values,
        theta=categories,
        fill='toself',
        name='Điểm số',
        line_color='#E74C3C'
    ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 10]
            )),
        showlegend=True,
        title="Đánh giá Hiệu suất Nhân viên",
        font=dict(size=12)
    )
    
    return fig

def create_360_performance_dashboard(analysis_360):
    """Create comprehensive 360-degree performance dashboard"""
    
    # Create subplots
    fig = make_subplots(
        rows=2, cols=3,
        subplot_titles=(
            'Điểm số 360',
            'So sánh với Đồng nghiệp',
            'Mạng lưới Hợp tác',
            'Phân bố Kỹ năng',
            'Xu hướng Phát triển',
            'Đánh giá Rủi ro'
        ),
        specs=[[{"type": "bar"}, {"type": "bar"}, {"type": "scatter"}],
               [{"type": "pie"}, {"type": "bar"}, {"type": "bar"}]]
    )
    
    # 360 Scores
    scores_360 = [
        analysis_360['collaboration_analysis']['networking_score'],
        analysis_360['skill_analysis']['expertise_level'],
        analysis_360['leadership_influence']['leadership_score'],
        analysis_360['cross_functional_contribution']['cross_functional_score'],
        analysis_360['team_impact']['team_impact_score'],
        analysis_360['innovation_contribution']['innovation_score']
    ]
    
    score_labels = [
        'Networking',
        'Chuyên môn',
        'Lãnh đạo',
        'Đa chức năng',
        'Tác động nhóm',
        'Đổi mới'
    ]
    
    fig.add_trace(
        go.Bar(x=score_labels, y=scores_360, name='Điểm 360'),
        row=1, col=1
    )
    
    # Peer comparison
    if 'peer_comparison' in analysis_360 and isinstance(analysis_360['performance_comparison'], dict):
        peer_data = analysis_360['performance_comparison']
        if 'completion_rate_percentile' in peer_data:
            fig.add_trace(
                go.Bar(
                    x=['Hoàn thành công việc', 'Số lượng dự án'],
                    y=[peer_data.get('completion_rate_percentile', 50), 
                       peer_data.get('project_count_percentile', 50)],
                    name='Percentile so với đồng nghiệp'
                ),
                row=1, col=2
            )
    
    # Collaboration network (simplified)
    collab_data = analysis_360['collaboration_analysis']
    fig.add_trace(
        go.Scatter(
            x=[0, collab_data['total_collaborators']],
            y=[0, collab_data['total_interactions']],
            mode='markers+text',
            text=['Nhân viên', 'Mạng lưới'],
            textposition="middle center",
            marker=dict(size=[30, collab_data['total_collaborators']*5]),
            name='Hợp tác'
        ),
        row=1, col=3
    )
    
    # Skills distribution
    skill_data = analysis_360['skill_analysis']
    if skill_data['top_skills']:
        skill_names = [skill[0][:20] for skill in skill_data['top_skills'][:5]]  # Truncate long names
        skill_counts = [skill[1] for skill in skill_data['top_skills'][:5]]
        
        fig.add_trace(
            go.Pie(
                labels=skill_names,
                values=skill_counts,
                name='Kỹ năng'
            ),
            row=2, col=1
        )
    
    # Growth trajectory
    growth_data = analysis_360['growth_trajectory']
    fig.add_trace(
        go.Bar(
            x=['Trajectory Score'],
            y=[growth_data['trajectory_score']],
            name='Phát triển'
        ),
        row=2, col=2
    )
    
    # Risk assessment
    risk_data = analysis_360['risk_assessment']
    risk_score = risk_data['risk_score']
    fig.add_trace(
        go.Bar(
            x=['Rủi ro'],
            y=[10 - risk_score],  # Invert so higher is better
            name='An toàn (10-Risk)',
            marker_color='red' if risk_score > 5 else 'orange' if risk_score > 2 else 'green'
        ),
        row=2, col=3
    )
    
    fig.update_layout(
        height=800,
        title_text="Dashboard 360 - Tổng quan Toàn diện",
        showlegend=False
    )
    
    return fig

def create_collaboration_network_chart(analysis_360):
    """Create collaboration network visualization"""
    collab_data = analysis_360['collaboration_analysis']
    
    if not collab_data['frequent_collaborators']:
        return None
    
    # Prepare data for network visualization
    collaborators = collab_data['frequent_collaborators'][:10]  # Top 10
    names = [collab[0] for collab in collaborators]
    interactions = [collab[1] for collab in collaborators]
    
    fig = go.Figure()
    
    # Add collaboration bars
    fig.add_trace(go.Bar(
        x=names,
        y=interactions,
        text=interactions,
        textposition='auto',
        name='Số lần hợp tác',
        marker_color='lightblue'
    ))
    
    fig.update_layout(
        title="Mạng lưới Hợp tác - Top Collaborators",
        xaxis_title="Đồng nghiệp",
        yaxis_title="Số lần hợp tác",
        xaxis_tickangle=-45
    )
    
    return fig

def create_skill_matrix_chart(analysis_360):
    """Create skill matrix visualization"""
    skill_data = analysis_360['skill_analysis']
    
    if not skill_data['top_skills']:
        return None
    
    skills = skill_data['top_skills'][:15]  # Top 15 skills
    skill_names = [skill[0] for skill in skills]
    skill_counts = [skill[1] for skill in skills]
    
    fig = px.bar(
        x=skill_counts,
        y=skill_names,
        orientation='h',
        title="Ma trận Kỹ năng - Chuyên môn",
        labels={'x': 'Số lần sử dụng', 'y': 'Kỹ năng'},
        color=skill_counts,
        color_continuous_scale='Blues'
    )
    
    fig.update_layout(height=600)
    
    return fig

def create_leadership_influence_chart(analysis_360):
    """Create leadership and influence visualization"""
    leadership_data = analysis_360['leadership_influence']
    team_impact_data = analysis_360['team_impact']
    
    categories = ['Ảnh hưởng', 'Được tìm kiếm', 'Lãnh đạo', 'Tác động nhóm']
    values = [
        leadership_data['influence_reach'],
        leadership_data['sought_after_score'] / 10,  # Scale down
        leadership_data['leadership_score'],
        team_impact_data['team_impact_score']
    ]
    
    fig = go.Figure()
    
    fig.add_trace(go.Scatterpolar(
        r=values,
        theta=categories,
        fill='toself',
        name='Leadership & Influence',
        line_color='green'
    ))
    
    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, max(values) * 1.1]
            )),
        title="Phân tích Lãnh đạo & Ảnh hưởng"
    )
    
    return fig

def create_basic_performance_chart(analysis_result):
    """Create bar chart for basic performance analysis"""
    stats = analysis_result.get('overall_stats', {})
    
    metrics = ['Tỷ lệ hoàn thành đúng hạn', 'Tuân thủ deadline (của task hoàn thành)', 'Ghi chép kết quả']
    values = [
        stats.get('on_time_completion_rate', 0),
        stats.get('on_time_rate', 0),
        stats.get('result_documentation_rate', 0)
    ]
    
    fig = px.bar(
        x=metrics,
        y=values,
        title="Chỉ số Hiệu suất Cơ bản (%)",
        color=values,
        color_continuous_scale='RdYlGn'
    )
    
    fig.update_layout(
        yaxis_title="Tỷ lệ (%)",
        showlegend=False
    )
    
    return fig

def create_data_quality_chart(analysis_result):
    """Create data quality visualization"""
    projects_analysis = analysis_result.get('projects_analysis', [])
    
    if not projects_analysis:
        return None
    
    project_names = [p['project_name'] for p in projects_analysis]
    quality_scores = [p['quality_score'] for p in projects_analysis]
    
    fig = px.bar(
        x=quality_scores,
        y=project_names,
        orientation='h',
        title="Chất lượng Dữ liệu theo Dự án",
        color=quality_scores,
        color_continuous_scale='RdYlGn'
    )
    
    fig.update_layout(
        xaxis_title="Điểm chất lượng",
        yaxis_title="Dự án"
    )
    
    return fig

def create_project_distribution_chart(df):
    """Create project distribution chart"""
    if 'Dự án' not in df.columns:
        return None
    
    project_counts = df['Dự án'].value_counts()
    
    fig = px.bar(
        x=project_counts.values,
        y=project_counts.index,
        orientation='h',
        title="Phân bố Công việc theo Dự án",
        labels={'x': 'Số lượng công việc', 'y': 'Dự án'}
    )
    
    return fig

def create_timeline_chart(df):
    """Create timeline chart for task completion and future deadlines"""
    if ('Ngày hoàn thành' not in df.columns and 'Deadline' not in df.columns) or 'Dự án' not in df.columns:
        return None
    
    timeline_data = []
    
    # Add completed tasks
    completed_tasks = df[df['Ngày hoàn thành'].notna()].copy()
    if not completed_tasks.empty:
        try:
            completed_tasks['Ngày hoàn thành'] = pd.to_datetime(completed_tasks['Ngày hoàn thành'], errors='coerce')
            completed_tasks = completed_tasks[completed_tasks['Ngày hoàn thành'].notna()]
            
            for _, task in completed_tasks.iterrows():
                timeline_data.append({
                    'Date': task['Ngày hoàn thành'],
                    'Project': task['Dự án'],
                    'Task': task['Tên công việc'],
                    'Status': task['Trạng thái'],
                    'Type': 'Hoàn thành'
                })
        except Exception as e:
            st.error(f"Error processing completed tasks: {e}")
    
    # Add future tasks
    future_tasks = df[(df['Ngày hoàn thành'].isna()) & (df['Deadline'].notna()) & (df['Trạng thái'] != 'Thất bại')].copy()
    if not future_tasks.empty:
        try:
            future_tasks['Deadline'] = pd.to_datetime(future_tasks['Deadline'], errors='coerce')
            future_tasks = future_tasks[future_tasks['Deadline'].notna()]
            
            for _, task in future_tasks.iterrows():
                timeline_data.append({
                    'Date': task['Deadline'],
                    'Project': task['Dự án'],
                    'Task': task['Tên công việc'],
                    'Status': task['Trạng thái'],
                    'Type': 'Deadline tương lai'
                })
        except Exception as e:
            st.error(f"Error processing future tasks: {e}")
    
    if not timeline_data:
        return None
    
    timeline_df = pd.DataFrame(timeline_data)
    timeline_df['Date'] = pd.to_datetime(timeline_df['Date'], errors='coerce')
    timeline_df = timeline_df[timeline_df['Date'].notna()]
    
    if timeline_df.empty:
        return None
    
    timeline_df = timeline_df.sort_values('Date')
    
    fig = px.scatter(
        timeline_df,
        x='Date',
        y='Project',
        color='Type',
        title="Timeline Công việc (Hoàn thành & Deadline tương lai)",
        hover_data=['Task', 'Status'],
        color_discrete_map={
            'Hoàn thành': '#2ECC71',
            'Deadline tương lai': '#F39C12'
        }
    )
    
    try:
        today = pd.Timestamp.now()
        
        if not timeline_df.empty:
            min_date = timeline_df['Date'].min()
            max_date = timeline_df['Date'].max()
            
            date_range_start = min_date - pd.Timedelta(days=30)
            date_range_end = max_date + pd.Timedelta(days=30)
            
            if date_range_start <= today <= date_range_end:
                fig.add_vline(
                    x=today,
                    line_dash="dash",
                    line_color="red",
                    annotation_text="Hôm nay"
                )
    except Exception as e:
        st.warning(f"Không thể thêm đường 'Hôm nay': {e}")
    
    return fig

def _create_project_indicators(stats: Dict, quality: Dict, risk: Dict) -> List[go.Indicator]:
    """Create indicator traces for project dashboard"""
    indicators = []

    # Completion rate indicator
    indicators.append(go.Indicator(
        mode="number+gauge+delta",
        value=stats['completion_rate'],
        title="Tỷ lệ Hoàn thành (%)",
        delta={'reference': 80},
        gauge={'axis': {'range': [0, 100]}},
        domain={'row': 0, 'column': 0}
    ))

    # Total tasks indicator
    indicators.append(go.Indicator(
        mode="number",
        value=stats['total_tasks'],
        title="Tổng số Task",
        domain={'row': 0, 'column': 1}
    ))

    # Quality indicator
    indicators.append(go.Indicator(
        mode="number+gauge",
        value=quality.get('description_rate', 0),
        title="Tỷ lệ có Mô tả (%)",
        gauge={'axis': {'range': [0, 100]}},
        domain={'row': 0, 'column': 2}
    ))

    # Risk indicator
    risk_colors = ['green', 'yellow', 'orange', 'red']
    risk_score = risk['risk_score']

    indicators.append(go.Indicator(
        mode="gauge+number",
        value=risk_score,
        title="Điểm Rủi ro",
        gauge={
            'axis': {'range': [0, 10]},
            'bar': {'color': risk_colors[min(3, risk_score // 3)]},
            'steps': [
                {'range': [0, 2], 'color': 'lightgreen'},
                {'range': [2, 4], 'color': 'yellow'},
                {'range': [4, 6], 'color': 'orange'},
                {'range': [6, 10], 'color': 'red'}
            ]
        },
        domain={'row': 2, 'column': 2}
    ))

    return indicators

def _create_task_distribution_charts(task_dist: Dict) -> List[go.Trace]:
    """Create task distribution visualization traces"""
    traces = []

    # Category pie chart
    if task_dist['by_category']:
        categories = list(task_dist['by_category'].keys())[:8]
        counts = list(task_dist['by_category'].values())[:8]

        traces.append(go.Pie(
            labels=categories,
            values=counts,
            name="Theo Danh mục",
            marker_colors=px.colors.qualitative.Set3
        ))
    else:
        traces.append(go.Pie(
            labels=['Chưa phân loại'],
            values=[1],
            name="Theo Danh mục"
        ))

    # Assignee bar chart
    if task_dist['by_assignee']:
        assignees = list(task_dist['by_assignee'].keys())[:10]
        counts = list(task_dist['by_assignee'].values())[:10]

        traces.append(go.Bar(
            x=counts,
            y=assignees,
            orientation='h',
            name="Theo Người thực hiện",
            marker_color='lightgreen'
        ))
    else:
        traces.append(go.Bar(
            x=[1],
            y=['Chưa phân công'],
            orientation='h',
            name="Theo Người thực hiện"
        ))

    return traces

def create_project_dashboard(project_analysis: Dict, employee_username: str = None) -> go.Figure:
    """Create comprehensive project dashboard"""
    stats = project_analysis['overall_stats']
    quality = project_analysis['quality_metrics']
    risk = project_analysis['risk_assessment']
    task_dist = project_analysis['task_analysis']

    # Create subplots
    fig = make_subplots(
        rows=3, cols=3,
        subplot_titles=(
            '📊 Tổng quan Dự án',
            '🎯 Chất lượng Dữ liệu',
            '⚠️ Đánh giá Rủi ro',
            '📋 Phân bố Task theo Danh mục',
            '👥 Phân bố theo Người thực hiện',
            '📈 Timeline & Deadline',
            '🤝 Mức độ Hợp tác',
            '📊 Tỷ lệ Hoàn thành',
            '💡 Khuyến nghị'
        ),
        specs=[
            [{"type": "indicator"}, {"type": "indicator"}, {"type": "indicator"}],
            [{"type": "pie"}, {"type": "bar"}, {"type": "scatter"}],
            [{"type": "bar"}, {"type": "bar"}, {"type": "domain"}]
        ]
    )

    # Add indicator traces
    indicators = _create_project_indicators(stats, quality, risk)
    for indicator in indicators:
        fig.add_trace(indicator)

    # Add task distribution traces
    dist_traces = _create_task_distribution_charts(task_dist)
    for i, trace in enumerate(dist_traces):
        row = 2 if i == 0 else 2
        col = 1 if i == 0 else 2
        fig.add_trace(trace, row=row, col=col)

    # Timeline visualization (simplified scatter)
    timeline = project_analysis['timeline_analysis']
    if timeline['tasks']:
        completed_tasks = [t for t in timeline['tasks'] if t['status'] == 'Hoàn thành' and t['completion_date']]
        if completed_tasks:
            fig.add_trace(
                go.Scatter(
                    x=[t['completion_date'] for t in completed_tasks[:20]],
                    y=[t['task_name'][:30] + "..." if len(t['task_name']) > 30 else t['task_name'] for t in completed_tasks[:20]],
                    mode='markers',
                    name='Đã hoàn thành',
                    marker=dict(color='green', size=8)
                ),
                row=2, col=3
            )

    # Collaboration metrics
    collab = project_analysis['collaboration_analysis']
    fig.add_trace(
        go.Bar(
            x=['Mật độ Hợp tác', 'Số Người Hợp tác'],
            y=[collab['collaboration_density'], collab['unique_collaborators']],
            name="Hợp tác",
            marker_color='lightblue'
        ),
        row=3, col=1
    )

    # Completion breakdown
    completion_data = [
        stats['completed_tasks'],
        stats['ongoing_tasks'],
        stats['failed_tasks']
    ]
    completion_labels = ['Hoàn thành', 'Đang thực hiện', 'Thất bại']

    fig.add_trace(
        go.Bar(
            x=completion_labels,
            y=completion_data,
            name="Trạng thái Task",
            marker_color=['green', 'orange', 'red']
        ),
        row=3, col=2
    )

    fig.update_layout(
        height=Config.CHART_HEIGHTS['dashboard'],
        title_text="📊 Dashboard Dự án Toàn diện",
        showlegend=False
    )

    return fig

def create_project_timeline_chart(project_analysis: Dict) -> go.Figure:
    """Create detailed timeline chart for project"""
    timeline = project_analysis['timeline_analysis']

    if not timeline['tasks']:
        return None

    # Prepare data for timeline
    tasks_data = []
    for task in timeline['tasks']:
        if task['start_date'] or task['deadline']:
            try:
                start_date = pd.to_datetime(task['start_date']) if task['start_date'] else None
                deadline_date = pd.to_datetime(task['deadline']) if task['deadline'] else None
                completion_date = pd.to_datetime(task['completion_date']) if task['completion_date'] else None

                if start_date or deadline_date:
                    tasks_data.append({
                        'Task': task['task_name'][:30] + "..." if len(task['task_name']) > 30 else task['task_name'],
                        'Start': start_date,
                        'Finish': deadline_date if deadline_date else (completion_date if completion_date else start_date),
                        'Status': task['status'],
                        'Resource': task.get('category', 'Chưa phân loại'),
                        'Overdue': task['is_overdue']
                    })
            except:
                continue

    if not tasks_data:
        return None

    df_timeline = pd.DataFrame(tasks_data)

    # Create Gantt-like chart
    fig = px.timeline(
        df_timeline,
        x_start="Start",
        x_end="Finish",
        y="Task",
        color="Status",
        color_discrete_map={
            "Hoàn thành": "green",
            "Đang thực hiện": "orange"
        },
        title="📅 Timeline Dự án",
        labels={'Status': 'Trạng thái'}
    )

    # Add vertical line for today
    today = datetime.now()
    fig.add_vline(
        x=today,
        line_dash="dash",
        line_color="red",
        annotation_text="Hôm nay"
    )

    # Highlight overdue tasks
    overdue_tasks = df_timeline[df_timeline['Overdue'] == True]
    if not overdue_tasks.empty:
        for _, task in overdue_tasks.iterrows():
            fig.add_vrect(
                x0=today,
                x1=task['Finish'],
                y0=task.name - 0.4,
                y1=task.name + 0.4,
                fillcolor="red",
                opacity=0.3,
                layer="below",
                line_width=0,
            )

    fig.update_layout(
        height=max(400, len(df_timeline) * 20),
        showlegend=True
    )

    return fig

def create_project_quality_radar(quality_metrics: Dict) -> go.Figure:
    """Create radar chart for project quality metrics"""
    categories = [
        'Mô tả chi tiết',
        'Deadline rõ ràng',
        'Kết quả đầy đủ',
        'Hợp tác tốt',
        'Chất lượng tổng thể'
    ]

    # Calculate quality scores (0-100)
    values = [
        quality_metrics.get('description_rate', 0),
        quality_metrics.get('deadline_rate', 0),
        quality_metrics.get('result_rate', 0),
        quality_metrics.get('follower_rate', 0),
        (quality_metrics.get('description_rate', 0) +
         quality_metrics.get('deadline_rate', 0) +
         quality_metrics.get('result_rate', 0) +
         quality_metrics.get('follower_rate', 0)) / 4
    ]

    fig = go.Figure()

    fig.add_trace(go.Scatterpolar(
        r=values,
        theta=categories,
        fill='toself',
        name='Chất lượng',
        line_color='#3498DB'
    ))

    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 100],
                tickformat=".0f"
            )),
        title="🎯 Chất lượng Dự án",
        showlegend=False
    )

    return fig

def create_employee_contribution_chart(project_analysis: Dict, employee_username: str) -> go.Figure:
    """Create chart showing employee's contribution to project"""
    employee_contrib = project_analysis.get('employee_contribution')

    if not employee_contrib or employee_contrib.get('message'):
        return None

    # Create metrics for employee
    metrics = [
        'Tỷ lệ hoàn thành',
        'Tỷ lệ đóng góp',
        'Số task được giao'
    ]

    values = [
        employee_contrib['completion_rate'],
        employee_contrib['contribution_percentage'],
        employee_contrib['total_tasks_assigned']
    ]

    fig = go.Figure()

    # Bar chart for metrics
    fig.add_trace(go.Bar(
        x=metrics,
        y=values,
        name='Đóng góp của Nhân viên',
        marker_color=['#2ECC71', '#3498DB', '#9B59B6']
    ))

    fig.update_layout(
        title=f"👤 Đóng góp của {employee_username} trong Dự án",
        yaxis_title="Giá trị",
        showlegend=False
    )

    # Add task categories breakdown
    if employee_contrib.get('task_categories'):
        categories = list(employee_contrib['task_categories'].keys())
        counts = list(employee_contrib['task_categories'].values())

        fig.add_trace(go.Pie(
            labels=categories,
            values=counts,
            name="Phân bố Task theo Danh mục",
            domain=dict(x=[0.7, 1.0], y=[0.0, 0.5])
        ))

    return fig

def create_project_risk_heatmap(risk_data: Dict) -> go.Figure:
    """Create risk assessment heatmap"""
    risks = risk_data.get('identified_risks', [])

    if not risks:
        # No risks - show positive message
        fig = go.Figure()
        fig.add_annotation(
            text="✅ Không có rủi ro đáng kể",
            xref="paper", yref="paper",
            x=0.5, y=0.5, showarrow=False,
            font=dict(size=20, color="green")
        )
        fig.update_layout(
            title="⚠️ Đánh giá Rủi ro Dự án",
            height=300
        )
        return fig

    # Create risk matrix
    risk_categories = []
    risk_levels = []
    risk_colors = []

    severity_map = {
        "Tỷ lệ hoàn thành thấp": 3,
        "Quá nhiều task quá hạn": 3,
        "Nhiều task thiếu mô tả": 2,
        "Thiếu sự hợp tác": 1
    }

    impact_map = {
        "Tỷ lệ hoàn thành thấp": 3,
        "Quá nhiều task quá hạn": 3,
        "Nhiều task thiếu mô tả": 2,
        "Thiếu sự hợp tác": 2
    }

    for risk in risks:
        risk_categories.append(risk)
        severity = severity_map.get(risk, 2)
        impact = impact_map.get(risk, 2)

        risk_levels.append(severity * impact)

        # Color based on risk level
        if severity * impact >= 6:
            risk_colors.append('red')
        elif severity * impact >= 4:
            risk_colors.append('orange')
        else:
            risk_colors.append('yellow')

    fig = go.Figure()

    fig.add_trace(go.Scatter(
        x=risk_categories,
        y=risk_levels,
        mode='markers+text',
        text=[f"Mức {level}" for level in risk_levels],
        textposition="top center",
        marker=dict(
            size=[level * 10 for level in risk_levels],
            color=risk_colors,
            showscale=True,
            colorscale='Reds',
            colorbar=dict(title="Mức độ Rủi ro")
        ),
        name="Rủi ro"
    ))

    fig.update_layout(
        title="⚠️ Ma trận Rủi ro Dự án",
        xaxis_title="Loại Rủi ro",
        yaxis_title="Mức độ Nghiêm trọng",
        height=400
    )

    return fig

# Main Streamlit App
def main():
    st.set_page_config(
        page_title="Hệ thống Phân tích Nhân viên với AI & Employee 360",
        page_icon="👥",
        layout="wide"
    )
    
    st.title("👥 Hệ thống Phân tích Hiệu suất Nhân viên với AI & Employee 360")
    st.markdown("### Đánh giá hiệu suất nhân viên theo thời gian bằng AI hoặc phân tích cơ bản + Phân tích 360 độ toàn diện")
    st.markdown("---")
    
    # Sidebar
    st.sidebar.header("🔧 Cấu hình")
    access_token = st.sidebar.text_input(
        "WeWork Access Token",
        value=WEWORK_ACCESS_TOKEN,
        type="password",
        help="Nhập access token của WeWork"
    )

    # Analysis mode selection
    st.sidebar.header("🎯 Chế độ Phân tích")
    analysis_mode = st.sidebar.selectbox(
        "Chọn phương pháp phân tích:",
        Config.ANALYSIS_MODES,
        help="Chọn phương pháp phân tích hiệu suất"
    )

    # Feature selection
    st.sidebar.header("🚀 Tính năng")
    feature_mode = st.sidebar.selectbox(
        "Chọn tính năng:",
        Config.FEATURE_MODES,
        help="Chọn loại phân tích muốn thực hiện"
    )

    # Show AI status
    if GENAI_AVAILABLE:
        if analysis_mode in ["AI", "Cả hai"]:
            st.sidebar.success("🤖 AI đã sẵn sàng với API key được cấu hình")
    else:
        if analysis_mode in ["AI", "Cả hai"]:
            st.sidebar.warning("⚠️ Cài đặt google-genai để sử dụng AI")

    if not access_token:
        st.warning("⚠️ Vui lòng nhập access token để tiếp tục")
        return

    # Initialize clients
    api_client = APIClient(access_token, ACCOUNT_ACCESS_TOKEN)

    # Initialize analyzers
    basic_analyzer = BasicAnalyzer()

    ai_analyzer = None
    if analysis_mode in ["AI", "Cả hai"] and GENAI_AVAILABLE:
        ai_analyzer = AIAnalyzer()
    elif analysis_mode in ["AI", "Cả hai"] and not GENAI_AVAILABLE:
        st.sidebar.warning("⚠️ Cài đặt google-genai để sử dụng AI")
    
    # Load employees
    if 'employees' not in st.session_state:
        with st.spinner("Đang tải danh sách nhân viên..."):
            try:
                st.session_state.employees = api_client.get_filtered_members()
            except Exception as e:
                st.error(f"❌ Lỗi khi tải danh sách nhân viên: {e}")
                return
    
    employees_df = st.session_state.employees
    
    if employees_df.empty:
        st.error("❌ Không có nhân viên nào được tìm thấy")
        return
    
    st.success(f"✅ Đã tìm thấy {len(employees_df)} nhân viên")
    
    # Route to appropriate interface based on feature selection
    if feature_mode == "Phân tích theo thời gian":
        employee_analysis_interface(api_client, basic_analyzer, ai_analyzer, analysis_mode)
    elif feature_mode == "Employee 360":
        employee_360_interface(api_client, employees_df, analysis_mode)
    elif feature_mode == "Phân tích Dự án":
        project_analysis_interface(api_client, employees_df)
    else:  # "Cả ba"
        tab1, tab2, tab3 = st.tabs(["⏰ Phân tích theo Thời gian", "🌟 Employee 360", "📊 Phân tích Dự án"])

        with tab1:
            employee_analysis_interface(api_client, basic_analyzer, ai_analyzer, analysis_mode)

        with tab2:
            employee_360_interface(api_client, employees_df, analysis_mode)

        with tab3:
            project_analysis_interface(api_client, employees_df)

def employee_360_interface(api_client, employees_df, analysis_mode):
    """Interface for Employee 360 analysis"""
    
    st.header("🌟 Employee 360 - Phân tích Toàn diện")
    st.markdown("Đánh giá 360 độ về hiệu suất, kỹ năng, hợp tác và ảnh hưởng của nhân viên")
    
    # Initialize 360 analyzer
    if 'employee_360_analyzer' not in st.session_state:
        st.session_state.employee_360_analyzer = Employee360Analyzer(api_client, employees_df)
    
    analyzer_360 = st.session_state.employee_360_analyzer
    
    # Load organization data button
    col1, col2 = st.columns([3, 1])
    
    with col2:
        if st.button("🔄 Tải dữ liệu tổ chức", type="secondary"):
            analyzer_360.load_organization_data()
            st.success("✅ Đã tải dữ liệu tổ chức!")
    
    with col1:
        if not analyzer_360.all_projects_data:
            st.warning("⚠️ Cần tải dữ liệu tổ chức trước khi thực hiện phân tích 360")
            return
        else:
            st.info(f"📊 Đã tải {len(analyzer_360.all_projects_data)} dự án cho phân tích 360")
    
    # Employee selection for 360 analysis
    st.subheader("👤 Chọn Nhân viên cho Phân tích 360")
    
    employee_options = {}
    for _, employee in employees_df.iterrows():
        display_name = f"{employee['name']} ({employee['username']}) - {employee['job']}"
        employee_options[display_name] = employee.to_dict()
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        selected_employee_display = st.selectbox(
            "Chọn nhân viên để phân tích 360:",
            options=list(employee_options.keys()),
            help="Chọn nhân viên từ danh sách để thực hiện phân tích toàn diện"
        )
    
    with col2:
        time_filter_360 = st.selectbox(
            "Khoảng thời gian:",
            ["Tất cả", "Năm hiện tại", "6 tháng gần nhất"],
            help="Chọn khoảng thời gian cho phân tích"
        )
    
    if selected_employee_display:
        selected_employee = employee_options[selected_employee_display]
        
        # Display employee info
        with st.expander("ℹ️ Thông tin Nhân viên", expanded=False):
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.markdown(f"**Tên:** {selected_employee['name']}")
            with col2:
                st.markdown(f"**Username:** {selected_employee['username']}")
            with col3:
                st.markdown(f"**Vị trí:** {selected_employee['job']}")
            with col4:
                st.markdown(f"**Email:** {selected_employee['email']}")
        
        # Analysis button
        if st.button("🚀 Thực hiện Phân tích 360", type="primary"):
            username = selected_employee['username']
            
            with st.spinner("🔍 Đang thực hiện phân tích 360 độ toàn diện..."):
                analysis_360 = analyzer_360.analyze_employee_360(username, time_filter_360)
                
                if analysis_360:
                    st.session_state.employee_360_data = {
                        'analysis_360': analysis_360,
                        'employee_info': selected_employee,
                        'time_filter': time_filter_360
                    }
                    st.success("✅ Phân tích 360 hoàn thành!")
                else:
                    st.error("❌ Không thể thực hiện phân tích 360")
    
    # Display 360 results
    display_employee_360_results()

def project_analysis_interface(api_client, employees_df):
    """Interface for Project Analysis"""

    st.header("📊 Phân tích Dự án Chi tiết")
    st.markdown("Phân tích toàn diện từng dự án với các metrics, timeline, chất lượng và rủi ro")

    # Initialize Project Analyzer
    if 'project_analyzer' not in st.session_state:
        st.session_state.project_analyzer = ProjectAnalyzer(api_client)

    project_analyzer = st.session_state.project_analyzer

    # Load all projects
    if 'all_projects' not in st.session_state:
        with st.spinner("Đang tải danh sách dự án..."):
            try:
                projects_response = requests.post(
                    "https://wework.base.vn/extapi/v3/project/list",
                    data={'access_token': api_client.goal_token},
                    timeout=REQUEST_TIMEOUT
                )
                projects_response.raise_for_status()
                all_projects = projects_response.json().get('projects', [])
                st.session_state.all_projects = all_projects
            except Exception as e:
                st.error(f"❌ Lỗi khi tải danh sách dự án: {e}")
                return

    all_projects = st.session_state.all_projects

    if not all_projects:
        st.error("❌ Không tìm thấy dự án nào")
        return

    st.success(f"✅ Đã tải {len(all_projects)} dự án")

    # Project selection
    st.subheader("🏗️ Chọn Dự án để Phân tích")

    project_options = {}
    for project in all_projects:
        display_name = f"{project['name']} (ID: {project['id']})"
        project_options[display_name] = project

    selected_project_display = st.selectbox(
        "Chọn dự án:",
        options=list(project_options.keys()),
        help="Chọn dự án từ danh sách để thực hiện phân tích chi tiết"
    )

    col1, col2 = st.columns([3, 1])

    with col1:
        include_employee_analysis = st.checkbox(
            "Phân tích đóng góp của nhân viên cụ thể",
            help="Chọn để phân tích đóng góp của một nhân viên trong dự án"
        )

    with col2:
        if st.button("🔄 Phân tích Dự án", type="primary"):
            if selected_project_display:
                selected_project = project_options[selected_project_display]
                employee_username = None

                if include_employee_analysis:
                    # Employee selection for contribution analysis
                    employee_options = {}
                    for _, employee in employees_df.iterrows():
                        display_name = f"{employee['name']} ({employee['username']})"
                        employee_options[display_name] = employee['username']

                    selected_employee_display = st.selectbox(
                        "Chọn nhân viên để phân tích đóng góp:",
                        options=list(employee_options.keys()),
                        key="employee_for_project"
                    )

                    if selected_employee_display:
                        employee_username = employee_options[selected_employee_display]

                # Perform project analysis
                with st.spinner("🔍 Đang phân tích dự án chi tiết..."):
                    try:
                        # Get full project data
                        project_detail_response = requests.post(
                            "https://wework.base.vn/extapi/v3/project/get.full",
                            data={
                                'access_token': api_client.goal_token,
                                'id': selected_project['id']
                            },
                            timeout=REQUEST_TIMEOUT
                        )
                        project_detail_response.raise_for_status()
                        project_data = project_detail_response.json()

                        # Perform analysis
                        project_analysis = project_analyzer.analyze_single_project(
                            project_data,
                            employee_username
                        )

                        st.session_state.project_analysis_data = {
                            'project_analysis': project_analysis,
                            'selected_project': selected_project,
                            'employee_username': employee_username
                        }

                        st.success("✅ Phân tích dự án hoàn thành!")

                    except Exception as e:
                        st.error(f"❌ Lỗi khi phân tích dự án: {str(e)}")

    # Display project analysis results
    display_project_analysis_results()

def display_project_analysis_results():
    """Display project analysis results"""
    if 'project_analysis_data' not in st.session_state:
        return

    data = st.session_state.project_analysis_data
    project_analysis = data['project_analysis']
    selected_project = data['selected_project']
    employee_username = data.get('employee_username')

    st.markdown("---")
    st.header(f"📊 Kết quả Phân tích: {selected_project['name']}")

    # Project overview
    st.subheader("🏗️ Thông tin Dự án")

    project_info = project_analysis['project_info']

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("ID Dự án", project_info['id'])
        st.metric("Trạng thái", project_info['status'])

    with col2:
        st.metric("Ưu tiên", project_info['priority'])
        st.metric("Tiến độ", f"{project_info['progress']}%")

    with col3:
        if project_info['start_time']:
            st.metric("Ngày bắt đầu", datetime.fromtimestamp(int(project_info['start_time'])).strftime('%Y-%m-%d'))
        if project_info['deadline']:
            st.metric("Deadline", datetime.fromtimestamp(int(project_info['deadline'])).strftime('%Y-%m-%d'))

    with col4:
        if project_info['description']:
            with st.expander("📝 Mô tả Dự án"):
                st.write(project_info['description'])

    # Main dashboard
    st.subheader("📈 Dashboard Tổng quan")

    dashboard_fig = create_project_dashboard(project_analysis, employee_username)
    if dashboard_fig:
        st.plotly_chart(dashboard_fig, use_container_width=True)

    # Detailed analysis tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📊 Thống kê Chi tiết",
        "📅 Timeline & Deadline",
        "🎯 Chất lượng & Hợp tác",
        "⚠️ Rủi ro & Khuyến nghị",
        "👤 Đóng góp Nhân viên" if employee_username else "📋 Phân tích Task"
    ])

    with tab1:
        display_project_detailed_stats(project_analysis)

    with tab2:
        display_project_timeline_analysis(project_analysis)

    with tab3:
        display_project_quality_collaboration(project_analysis)

    with tab4:
        display_project_risk_recommendations(project_analysis)

    with tab5:
        if employee_username:
            display_employee_contribution_in_project(project_analysis, employee_username)
        else:
            display_task_distribution_analysis(project_analysis)

    # Export section
    st.markdown("---")
    st.subheader("📤 Xuất Báo cáo Dự án")

    col1, col2, col3 = st.columns(3)

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        if st.button("📄 Tải báo cáo Word", type="primary"):
            try:
                doc = export_project_to_word(project_analysis, selected_project, employee_username)
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)

                st.download_button(
                    label="📥 Tải xuống Word",
                    data=doc_buffer.getvalue(),
                    file_name=f"project_{selected_project['id']}_analysis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.success("✅ Báo cáo Word đã được tạo!")

            except Exception as e:
                st.error(f"❌ Lỗi khi tạo báo cáo Word: {str(e)}")

    with col2:
        # Export as JSON
        export_data = {
            'project_info': selected_project,
            'analysis': project_analysis,
            'employee_focus': employee_username
        }

        json_data = json.dumps(export_data, ensure_ascii=False, indent=2)
        st.download_button(
            label="📋 Tải dữ liệu JSON",
            data=json_data,
            file_name=f"project_{selected_project['id']}_analysis.json",
            mime="application/json"
        )

    with col3:
        # Export summary CSV
        summary_data = create_project_summary_csv(project_analysis)
        if summary_data:
            csv_data = summary_data.to_csv(index=False, encoding='utf-8-sig')
            st.download_button(
                label="📊 Tải tóm tắt CSV",
                data=csv_data,
                file_name=f"project_{selected_project['id']}_summary.csv",
                mime="text/csv"
            )

    with col4:
        # Export timeline Excel
        if st.button("📊 Tải Timeline Excel"):
            excel_data = export_project_timeline_to_excel(project_analysis, selected_project)
            if excel_data:
                st.download_button(
                    label="📥 Tải xuống Excel",
                    data=excel_data,
                    file_name=f"project_{selected_project['id']}_timeline.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
                st.success("✅ File Excel đã được tạo!")
            else:
                st.error("❌ Không thể tạo file Excel")

        # PDF export
        if st.button("📄 Tải báo cáo PDF"):
            pdf_data = create_project_pdf_report(project_analysis, selected_project, employee_username)
            if pdf_data:
                st.download_button(
                    label="📥 Tải xuống PDF",
                    data=pdf_data,
                    file_name=f"project_{selected_project['id']}_report.pdf",
                    mime="application/pdf"
                )
                st.success("✅ Báo cáo PDF đã được tạo!")
            else:
                st.info("💡 Cần cài đặt reportlab để xuất PDF: pip install reportlab")

def display_project_detailed_stats(project_analysis):
    """Display detailed project statistics"""
    stats = project_analysis['overall_stats']

    col1, col2, col3, col4 = st.columns(4)

    with col1:
        st.metric("📊 Tổng Task", stats['total_tasks'])
        st.metric("✅ Hoàn thành", stats['completed_tasks'])
        st.metric("🔄 Đang thực hiện", stats['ongoing_tasks'])

    with col2:
        st.metric("❌ Thất bại", stats['failed_tasks'])
        st.metric("⏰ Đúng hạn", stats['on_time_tasks'])
        st.metric("⚠️ Quá hạn", stats['overdue_tasks'])

    with col3:
        st.metric("📈 Tỷ lệ Hoàn thành", f"{stats['completion_rate']}%")
        st.metric("⏱️ Tỷ lệ Đúng hạn", f"{stats['on_time_rate']}%")
        st.metric("📊 TB Hoàn thành", f"{stats['avg_completion_percentage']}%")

    with col4:
        # Progress visualization
        progress_color = "green" if stats['completion_rate'] > 80 else "orange" if stats['completion_rate'] > 50 else "red"
        st.markdown(f"""
        <div style="background-color: #f0f0f0; border-radius: 10px; padding: 10px; margin: 5px;">
            <div style="background-color: {progress_color}; width: {stats['completion_rate']}%; height: 20px; border-radius: 10px;"></div>
            <p style="text-align: center; margin: 5px 0 0 0; color: {progress_color}; font-weight: bold;">
                Tiến độ: {stats['completion_rate']}%
            </p>
        </div>
        """, unsafe_allow_html=True)

def display_project_timeline_analysis(project_analysis):
    """Display project timeline analysis"""
    timeline = project_analysis['timeline_analysis']

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 📅 Deadline Sắp tới")
        upcoming = timeline.get('upcoming_deadlines', [])

        if upcoming:
            for deadline in upcoming:
                days_text = "hôm nay" if deadline['days_until'] == 0 else f"{deadline['days_until']} ngày tới"
                st.markdown(f"• **{deadline['task_name']}** - {days_text}")
        else:
            st.info("✅ Không có deadline nào trong 7 ngày tới")

    with col2:
        st.markdown("### ⚠️ Task Quá hạn")
        overdue = timeline.get('overdue_tasks', [])

        if overdue:
            for task in overdue:
                st.error(f"• {task['task_name']} - Quá hạn {abs((datetime.now() - datetime.strptime(task['deadline'], '%Y-%m-%d')).days)} ngày")
        else:
            st.success("✅ Không có task nào quá hạn")

    # Timeline chart
    timeline_fig = create_project_timeline_chart(project_analysis)
    if timeline_fig:
        st.plotly_chart(timeline_fig, use_container_width=True)

def display_project_quality_collaboration(project_analysis):
    """Display quality and collaboration analysis"""
    quality = project_analysis['quality_metrics']
    collaboration = project_analysis['collaboration_analysis']

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 🎯 Chất lượng Dữ liệu")

        quality_fig = create_project_quality_radar(quality)
        if quality_fig:
            st.plotly_chart(quality_fig, use_container_width=True)

        st.markdown(f"**Tỷ lệ có mô tả:** {quality.get('description_rate', 0)}%")
        st.markdown(f"**Tỷ lệ có deadline:** {quality.get('deadline_rate', 0)}%")
        st.markdown(f"**Tỷ lệ có kết quả:** {quality.get('result_rate', 0)}%")
        st.markdown(f"**Tỷ lệ có người theo dõi:** {quality.get('follower_rate', 0)}%")

    with col2:
        st.markdown("### 🤝 Phân tích Hợp tác")

        st.metric("👥 Số người hợp tác", collaboration['unique_collaborators'])
        st.metric("🔗 Tổng tương tác", collaboration['total_collaborations'])
        st.metric("📊 Mật độ hợp tác", f"{collaboration['collaboration_density']}/task")

        # Collaboration visualization
        if collaboration['collaboration_matrix']:
            # Simple bar chart for top collaborators
            collab_data = []
            for user, data in collaboration['collaboration_matrix'].items():
                collab_data.append({
                    'Người dùng': user,
                    'Số tương tác': len(data['followers']) + len(data['following'])
                })

            collab_df = pd.DataFrame(collab_data).nlargest(10, 'Số tương tác')

            fig = px.bar(
                collab_df,
                x='Số tương tác',
                y='Người dùng',
                orientation='h',
                title="🏆 Top Collaborators",
                color='Số tương tác',
                color_continuous_scale='Blues'
            )
            st.plotly_chart(fig, use_container_width=True)

def display_project_risk_recommendations(project_analysis):
    """Display risk assessment and recommendations"""
    risk = project_analysis['risk_assessment']
    recommendations = project_analysis['recommendations']

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### ⚠️ Đánh giá Rủi ro")

        risk_fig = create_project_risk_heatmap(risk)
        st.plotly_chart(risk_fig, use_container_width=True)

        st.metric("🏆 Mức độ Rủi ro", risk['risk_level'])
        st.metric("📊 Điểm Rủi ro", f"{risk['risk_score']}/10")

    with col2:
        st.markdown("### 💡 Khuyến nghị Cải thiện")

        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                st.markdown(f"{i}. {rec}")
        else:
            st.success("✅ Dự án đang tiến triển tốt!")

def display_employee_contribution_in_project(project_analysis, employee_username):
    """Display employee's contribution in the project"""
    employee_contrib = project_analysis.get('employee_contribution')

    if not employee_contrib or employee_contrib.get('message'):
        st.warning(f"⚠️ Không tìm thấy thông tin đóng góp của {employee_username} trong dự án này")
        return

    st.markdown(f"### 👤 Đóng góp của {employee_username}")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric("📋 Task được giao", employee_contrib['total_tasks_assigned'])
        st.metric("✅ Đã hoàn thành", employee_contrib['completed_tasks'])

    with col2:
        st.metric("📈 Tỷ lệ hoàn thành", f"{employee_contrib['completion_rate']}%")
        st.metric("🏆 Tỷ lệ đóng góp", f"{employee_contrib['contribution_percentage']}%")

    with col3:
        # Task categories
        if employee_contrib.get('task_categories'):
            categories = list(employee_contrib['task_categories'].keys())
            counts = list(employee_contrib['task_categories'].values())

            fig = px.pie(
                values=counts,
                names=categories,
                title="📊 Phân bố Task theo Danh mục"
            )
            st.plotly_chart(fig, use_container_width=True)

    # Employee contribution chart
    contrib_fig = create_employee_contribution_chart(project_analysis, employee_username)
    if contrib_fig:
        st.plotly_chart(contrib_fig, use_container_width=True)

    # Task details
    if employee_contrib.get('task_details'):
        st.markdown("### 📋 Chi tiết Task")

        task_df = pd.DataFrame(employee_contrib['task_details'])
        st.dataframe(task_df, use_container_width=True)

def display_task_distribution_analysis(project_analysis):
    """Display task distribution analysis"""
    task_dist = project_analysis['task_analysis']

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 📋 Phân bố theo Danh mục")

        if task_dist['by_category']:
            category_df = pd.DataFrame([
                {'Danh mục': cat, 'Số lượng': count}
                for cat, count in task_dist['by_category'].items()
            ])

            fig = px.bar(
                category_df,
                x='Số lượng',
                y='Danh mục',
                orientation='h',
                title="Task theo Danh mục",
                color='Số lượng',
                color_continuous_scale='Viridis'
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Không có dữ liệu phân loại task")

    with col2:
        st.markdown("### 👥 Phân bố theo Người thực hiện")

        if task_dist['by_assignee']:
            assignee_df = pd.DataFrame([
                {'Người thực hiện': assignee, 'Số lượng': count}
                for assignee, count in task_dist['by_assignee'].items()
            ]).nlargest(15, 'Số lượng')

            fig = px.bar(
                assignee_df,
                x='Số lượng',
                y='Người thực hiện',
                orientation='h',
                title="Task theo Người thực hiện",
                color='Số lượng',
                color_continuous_scale='Plasma'
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("Không có dữ liệu phân công task")

def create_project_summary_csv(project_analysis):
    """Create summary CSV for project analysis"""
    try:
        stats = project_analysis['overall_stats']
        quality = project_analysis['quality_metrics']
        risk = project_analysis['risk_assessment']

        summary_data = {
            'Metric': [
                'Tổng Task', 'Task Hoàn thành', 'Task Đang thực hiện', 'Task Thất bại',
                'Tỷ lệ Hoàn thành (%)', 'Tỷ lệ Đúng hạn (%)', 'Task Quá hạn',
                'Tỷ lệ có Mô tả (%)', 'Tỷ lệ có Deadline (%)', 'Tỷ lệ có Kết quả (%)',
                'Mật độ Hợp tác', 'Điểm Rủi ro', 'Mức độ Rủi ro'
            ],
            'Value': [
                stats['total_tasks'], stats['completed_tasks'], stats['ongoing_tasks'], stats['failed_tasks'],
                stats['completion_rate'], stats['on_time_rate'], stats['overdue_tasks'],
                quality.get('description_rate', 0), quality.get('deadline_rate', 0), quality.get('result_rate', 0),
                project_analysis['collaboration_analysis']['collaboration_density'],
                risk['risk_score'], risk['risk_level']
            ]
        }

        return pd.DataFrame(summary_data)

    except Exception as e:
        st.error(f"❌ Lỗi tạo CSV tóm tắt: {e}")
        return None

def export_project_to_word(project_analysis, selected_project, employee_username=None):
    """Export project analysis to Word document"""
    doc = Document()

    # Title
    title = doc.add_heading(f'Báo cáo Phân tích Dự án: {selected_project["name"]}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project info
    doc.add_heading('🏗️ Thông tin Dự án', level=1)

    project_info = project_analysis['project_info']
    p = doc.add_paragraph()
    p.add_run('ID: ').bold = True
    p.add_run(project_info['id'])
    p.add_run(' | Trạng thái: ').bold = True
    p.add_run(project_info['status'])
    p.add_run(' | Ưu tiên: ').bold = True
    p.add_run(project_info['priority'])

    if project_info.get('description'):
        doc.add_heading('📝 Mô tả:', level=2)
        doc.add_paragraph(project_info['description'])

    # Overall stats
    doc.add_heading('📊 Thống kê Tổng quan', level=1)
    stats = project_analysis['overall_stats']

    stats_table = doc.add_table(rows=6, cols=2)
    stats_table.style = 'Table Grid'

    stats_data = [
        ('Tổng Task', stats['total_tasks']),
        ('Hoàn thành', stats['completed_tasks']),
        ('Đang thực hiện', stats['ongoing_tasks']),
        ('Thất bại', stats['failed_tasks']),
        ('Tỷ lệ Hoàn thành', f"{stats['completion_rate']}%"),
        ('Tỷ lệ Đúng hạn', f"{stats['on_time_rate']}%")
    ]

    for i, (metric, value) in enumerate(stats_data):
        stats_table.cell(i, 0).text = metric
        stats_table.cell(i, 1).text = str(value)

    # Quality metrics
    doc.add_heading('🎯 Chất lượng Dữ liệu', level=1)
    quality = project_analysis['quality_metrics']

    quality_table = doc.add_table(rows=4, cols=2)
    quality_table.style = 'Table Grid'

    quality_data = [
        ('Tỷ lệ có Mô tả', f"{quality.get('description_rate', 0)}%"),
        ('Tỷ lệ có Deadline', f"{quality.get('deadline_rate', 0)}%"),
        ('Tỷ lệ có Kết quả', f"{quality.get('result_rate', 0)}%"),
        ('Tỷ lệ có Người theo dõi', f"{quality.get('follower_rate', 0)}%")
    ]

    for i, (metric, value) in enumerate(quality_data):
        quality_table.cell(i, 0).text = metric
        quality_table.cell(i, 1).text = value

    # Task distribution
    task_dist = project_analysis['task_analysis']
    if task_dist['by_category']:
        doc.add_heading('📋 Phân bố Task theo Danh mục', level=1)
        for category, count in list(task_dist['by_category'].items())[:10]:
            p = doc.add_paragraph()
            p.add_run(f'{category}: ').bold = True
            p.add_run(f'{count} task')

    # Collaboration analysis
    doc.add_heading('🤝 Phân tích Hợp tác', level=1)
    collab = project_analysis['collaboration_analysis']

    p = doc.add_paragraph()
    p.add_run('Số người hợp tác: ').bold = True
    p.add_run(str(collab['unique_collaborators']))
    p.add_run(' | Tổng tương tác: ').bold = True
    p.add_run(str(collab['total_collaborations']))
    p.add_run(' | Mật độ: ').bold = True
    p.add_run(f"{collab['collaboration_density']}/task")

    # Risk assessment
    doc.add_heading('⚠️ Đánh giá Rủi ro', level=1)
    risk = project_analysis['risk_assessment']

    p = doc.add_paragraph()
    p.add_run('Mức độ Rủi ro: ').bold = True
    p.add_run(risk['risk_level'])
    p.add_run(' | Điểm Rủi ro: ').bold = True
    p.add_run(f"{risk['risk_score']}/10")

    if risk['identified_risks']:
        doc.add_heading('Rủi ro được xác định:', level=2)
        for risk_item in risk['identified_risks']:
            doc.add_paragraph(f'• {risk_item}')

    # Recommendations
    doc.add_heading('💡 Khuyến nghị', level=1)
    recommendations = project_analysis['recommendations']
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(f'{i}. {rec}')

    # Employee contribution (if applicable)
    if employee_username and project_analysis.get('employee_contribution'):
        doc.add_heading(f'👤 Đóng góp của {employee_username}', level=1)
        emp_contrib = project_analysis['employee_contribution']

        p = doc.add_paragraph()
        p.add_run('Task được giao: ').bold = True
        p.add_run(str(emp_contrib['total_tasks_assigned']))
        p.add_run(' | Hoàn thành: ').bold = True
        p.add_run(str(emp_contrib['completed_tasks']))
        p.add_run(' | Tỷ lệ: ').bold = True
        p.add_run(f"{emp_contrib['completion_rate']}%")
        p.add_run(' | Đóng góp: ').bold = True
        p.add_run(f"{emp_contrib['contribution_percentage']}%")

        if emp_contrib.get('task_categories'):
            doc.add_heading('Phân bố Task theo Danh mục:', level=2)
            for category, count in emp_contrib['task_categories'].items():
                doc.add_paragraph(f'• {category}: {count} task')

        # Employee task details
        if emp_contrib.get('task_details'):
            doc.add_heading('Chi tiết Task:', level=2)
            for task in emp_contrib['task_details'][:20]:  # Limit to prevent too large document
                p = doc.add_paragraph()
                p.add_run(f'• {task["name"]} - {task["status"]} ({task["completion_percentage"]}%)')
                if task.get('category'):
                    p.add_run(f' | Danh mục: {task["category"]}')

    # Add timestamp
    doc.add_paragraph()
    doc.add_paragraph(f"Báo cáo được tạo vào: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    return doc

def export_project_timeline_to_excel(project_analysis, selected_project):
    """Export project timeline to Excel"""
    try:
        timeline = project_analysis['timeline_analysis']
        if not timeline['tasks']:
            return None

        # Create DataFrame for timeline
        timeline_data = []
        for task in timeline['tasks']:
            timeline_data.append({
                'Tên Task': task['task_name'],
                'Danh mục': task.get('category', 'Chưa phân loại'),
                'Trạng thái': task['status'],
                'Ngày bắt đầu': task['start_date'] or '',
                'Deadline': task['deadline'] or '',
                'Ngày hoàn thành': task['completion_date'] or '',
                'Quá hạn': 'Có' if task['is_overdue'] else 'Không'
            })

        df_timeline = pd.DataFrame(timeline_data)

        # Create Excel writer
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df_timeline.to_excel(writer, sheet_name='Timeline', index=False)

            # Add summary sheet
            summary_data = create_project_summary_csv(project_analysis)
            if summary_data is not None:
                summary_data.to_excel(writer, sheet_name='Tóm tắt', index=False)

        output.seek(0)
        return output.getvalue()

    except Exception as e:
        st.error(f"❌ Lỗi xuất Excel: {e}")
        return None

def create_project_pdf_report(project_analysis, selected_project, employee_username=None):
    """Create PDF report for project analysis"""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        story.append(Paragraph(f"Báo cáo Phân tích Dự án: {selected_project['name']}", title_style))
        story.append(Spacer(1, 12))

        # Project info
        story.append(Paragraph("🏗️ THÔNG TIN DỰ ÁN", styles['Heading1']))
        project_info = project_analysis['project_info']

        info_data = [
            ['ID Dự án', project_info['id']],
            ['Trạng thái', project_info['status']],
            ['Ưu tiên', project_info['priority']],
            ['Tiến độ', f"{project_info['progress']}%"]
        ]

        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(info_table)
        story.append(Spacer(1, 12))

        # Statistics
        story.append(Paragraph("📊 THỐNG KÊ TỔNG QUAN", styles['Heading1']))
        stats = project_analysis['overall_stats']

        stats_data = [
            ['Tổng Task', str(stats['total_tasks'])],
            ['Hoàn thành', str(stats['completed_tasks'])],
            ['Đang thực hiện', str(stats['ongoing_tasks'])],
            ['Thất bại', str(stats['failed_tasks'])],
            ['Tỷ lệ Hoàn thành', f"{stats['completion_rate']}%"],
            ['Tỷ lệ Đúng hạn', f"{stats['on_time_rate']}%"]
        ]

        stats_table = Table(stats_data, colWidths=[3*inch, 3*inch])
        stats_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(stats_table)
        story.append(Spacer(1, 12))

        # Risk assessment
        story.append(Paragraph("⚠️ ĐÁNH GIÁ RỦI RO", styles['Heading1']))
        risk = project_analysis['risk_assessment']

        risk_text = f"Mức độ Rủi ro: {risk['risk_level']} | Điểm Rủi ro: {risk['risk_score']}/10"
        story.append(Paragraph(risk_text, styles['Normal']))
        story.append(Spacer(1, 6))

        if risk['identified_risks']:
            story.append(Paragraph("Rủi ro được xác định:", styles['Heading2']))
            for risk_item in risk['identified_risks']:
                story.append(Paragraph(f"• {risk_item}", styles['Normal']))
        story.append(Spacer(1, 12))

        # Recommendations
        story.append(Paragraph("💡 KHUYẾN NGHỊ", styles['Heading1']))
        recommendations = project_analysis['recommendations']
        for i, rec in enumerate(recommendations, 1):
            story.append(Paragraph(f"{i}. {rec}", styles['Normal']))

        # Employee contribution (if applicable)
        if employee_username and project_analysis.get('employee_contribution'):
            story.append(PageBreak())
            story.append(Paragraph(f"👤 ĐÓNG GÓP CỦA {employee_username.upper()}", styles['Heading1']))
            emp_contrib = project_analysis['employee_contribution']

            contrib_data = [
                ['Task được giao', str(emp_contrib['total_tasks_assigned'])],
                ['Hoàn thành', str(emp_contrib['completed_tasks'])],
                ['Tỷ lệ hoàn thành', f"{emp_contrib['completion_rate']}%"],
                ['Tỷ lệ đóng góp', f"{emp_contrib['contribution_percentage']}%"]
            ]

            contrib_table = Table(contrib_data, colWidths=[3*inch, 3*inch])
            contrib_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(contrib_table)

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    except ImportError:
        st.error("❌ Cần cài đặt reportlab để xuất PDF: pip install reportlab")
        return None
    except Exception as e:
        st.error(f"❌ Lỗi tạo PDF: {e}")
        return None

def display_employee_360_results():
    """Display Employee 360 analysis results"""
    if 'employee_360_data' not in st.session_state:
        return
    
    data = st.session_state.employee_360_data
    analysis_360 = data['analysis_360']
    employee_info = data['employee_info']
    time_filter = data['time_filter']
    
    st.markdown("---")
    st.header(f"🌟 Kết quả Phân tích 360: {employee_info['name']}")
    
    # Overview metrics
    st.subheader("📊 Tổng quan 360")
    
    col1, col2, col3, col4, col5, col6 = st.columns(6)
    
    with col1:
        st.metric(
            "Networking", 
            f"{analysis_360['collaboration_analysis']['networking_score']:.1f}/10",
            help="Điểm mạng lưới hợp tác"
        )
    
    with col2:
        st.metric(
            "Chuyên môn", 
            f"{analysis_360['skill_analysis']['expertise_level']:.1f}/10",
            help="Mức độ chuyên môn"
        )
    
    with col3:
        st.metric(
            "Lãnh đạo", 
            f"{analysis_360['leadership_influence']['leadership_score']:.1f}/10",
            help="Khả năng lãnh đạo"
        )
    
    with col4:
        st.metric(
            "Đa chức năng", 
            f"{analysis_360['cross_functional_contribution']['cross_functional_score']:.1f}/10",
            help="Khả năng làm việc đa chức năng"
        )
    
    with col5:
        st.metric(
            "Tác động", 
            f"{analysis_360['team_impact']['team_impact_score']:.1f}/10",
            help="Tác động đến nhóm"
        )
    
    with col6:
        st.metric(
            "Đổi mới", 
            f"{analysis_360['innovation_contribution']['innovation_score']:.1f}/10",
            help="Khả năng đổi mới"
        )
    
    # Main dashboard
    st.subheader("📈 Dashboard 360")
    dashboard_fig = create_360_performance_dashboard(analysis_360)
    st.plotly_chart(dashboard_fig, use_container_width=True)
    
    # Detailed analysis tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "🤝 Hợp tác & Mạng lưới", 
        "🎯 Kỹ năng & Chuyên môn", 
        "👑 Lãnh đạo & Ảnh hưởng", 
        "📈 Phát triển & Xu hướng", 
        "⚠️ Rủi ro & Khuyến nghị"
    ])
    
    with tab1:
        display_collaboration_analysis(analysis_360)
    
    with tab2:
        display_skill_analysis(analysis_360)
    
    with tab3:
        display_leadership_analysis(analysis_360)
    
    with tab4:
        display_growth_analysis(analysis_360)
    
    with tab5:
        display_risk_recommendations(analysis_360)
    
    # Export section for 360
    st.markdown("---")
    st.subheader("📤 Xuất Báo cáo 360")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📄 Tải báo cáo 360 Word", type="primary"):
            try:
                doc = export_360_to_word(analysis_360, employee_info, time_filter)
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                st.download_button(
                    label="📥 Tải xuống báo cáo 360",
                    data=doc_buffer.getvalue(),
                    file_name=f"{employee_info['username']}_360_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("✅ Báo cáo 360 Word đã được tạo!")
                
            except Exception as e:
                st.error(f"❌ Lỗi khi tạo báo cáo Word: {str(e)}")
    
    with col2:
        # Export analysis results as JSON
        export_data_360 = {
            'employee_info': employee_info,
            'time_filter': time_filter,
            'analysis_360': analysis_360,
            'summary': {
                'networking_score': analysis_360['collaboration_analysis']['networking_score'],
                'expertise_level': analysis_360['skill_analysis']['expertise_level'],
                'leadership_score': analysis_360['leadership_influence']['leadership_score'],
                'cross_functional_score': analysis_360['cross_functional_contribution']['cross_functional_score'],
                'team_impact_score': analysis_360['team_impact']['team_impact_score'],
                'innovation_score': analysis_360['innovation_contribution']['innovation_score']
            }
        }
        
        json_data_360 = json.dumps(export_data_360, ensure_ascii=False, indent=2)
        st.download_button(
            label="📋 Tải dữ liệu 360 JSON",
            data=json_data_360,
            file_name=f"{employee_info['username']}_360_analysis.json",
            mime="application/json"
        )
    
    with col3:
        # Create summary CSV
        summary_data = {
            'Metric': [
                'Networking Score', 'Expertise Level', 'Leadership Score',
                'Cross-functional Score', 'Team Impact Score', 'Innovation Score',
                'Total Collaborators', 'Skill Diversity', 'Growth Trend'
            ],
            'Value': [
                analysis_360['collaboration_analysis']['networking_score'],
                analysis_360['skill_analysis']['expertise_level'],
                analysis_360['leadership_influence']['leadership_score'],
                analysis_360['cross_functional_contribution']['cross_functional_score'],
                analysis_360['team_impact']['team_impact_score'],
                analysis_360['innovation_contribution']['innovation_score'],
                analysis_360['collaboration_analysis']['total_collaborators'],
                analysis_360['skill_analysis']['skill_diversity'],
                analysis_360['growth_trajectory']['trajectory_score']
            ]
        }
        
        summary_df = pd.DataFrame(summary_data)
        csv_data_360 = summary_df.to_csv(index=False, encoding='utf-8-sig')
        st.download_button(
            label="📊 Tải tóm tắt CSV",
            data=csv_data_360,
            file_name=f"{employee_info['username']}_360_summary.csv",
            mime="text/csv"
        )

def display_collaboration_analysis(analysis_360):
    """Display collaboration and networking analysis"""
    collab_data = analysis_360['collaboration_analysis']
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📊 Thống kê Hợp tác")
        st.markdown(f"**Tổng số đồng nghiệp hợp tác:** {collab_data['total_collaborators']}")
        st.markdown(f"**Tổng số tương tác:** {collab_data['total_interactions']}")
        st.markdown(f"**Điểm networking:** {collab_data['networking_score']:.1f}/10")
        st.markdown(f"**Mức độ đa dạng hợp tác:** {collab_data['collaboration_diversity']}")
    
    with col2:
        # Collaboration network chart
        collab_chart = create_collaboration_network_chart(analysis_360)
        if collab_chart:
            st.plotly_chart(collab_chart, use_container_width=True)
    
    # Top collaborators
    if collab_data['frequent_collaborators']:
        st.markdown("### 🤝 Top Collaborators")
        collab_df = pd.DataFrame(
            collab_data['frequent_collaborators'], 
            columns=['Đồng nghiệp', 'Số lần hợp tác']
        )
        st.dataframe(collab_df, use_container_width=True)

def display_skill_analysis(analysis_360):
    """Display skill and expertise analysis"""
    skill_data = analysis_360['skill_analysis']
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 🎯 Thống kê Kỹ năng")
        st.markdown(f"**Số lĩnh vực kỹ năng:** {skill_data['skill_diversity']}")
        st.markdown(f"**Mức độ chuyên môn:** {skill_data['expertise_level']:.1f}/10")
        st.markdown(f"**Tổng kinh nghiệm:** {skill_data['total_skill_instances']} lần sử dụng")
    
    with col2:
        # Skill matrix chart
        skill_chart = create_skill_matrix_chart(analysis_360)
        if skill_chart:
            st.plotly_chart(skill_chart, use_container_width=True)
    
    # Top skills
    if skill_data['top_skills']:
        st.markdown("### 🏆 Top Kỹ năng")
        skill_df = pd.DataFrame(
            skill_data['top_skills'][:10], 
            columns=['Kỹ năng', 'Số lần sử dụng']
        )
        st.dataframe(skill_df, use_container_width=True)

def display_leadership_analysis(analysis_360):
    """Display leadership and influence analysis"""
    leadership_data = analysis_360['leadership_influence']
    team_impact_data = analysis_360['team_impact']
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 👑 Phân tích Lãnh đạo")
        st.markdown(f"**Phạm vi ảnh hưởng:** {leadership_data['influence_reach']} người")
        st.markdown(f"**Được tìm kiếm:** {leadership_data['sought_after_score']} lần")
        st.markdown(f"**Điểm lãnh đạo:** {leadership_data['leadership_score']:.1f}/10")
        st.markdown(f"**Mức độ lãnh đạo:** {leadership_data['leadership_level']}")
        
        st.markdown("### 🎯 Tác động Nhóm")
        st.markdown(f"**Điểm tác động nhóm:** {team_impact_data['team_impact_score']:.1f}/10")
        st.markdown(f"**Phạm vi nhóm:** {team_impact_data['team_reach']} người")
        st.markdown(f"**Chia sẻ kiến thức:** {team_impact_data['knowledge_sharing_instances']} lần")
        st.markdown(f"**Lĩnh vực chuyên môn:** {team_impact_data['expertise_areas']}")
    
    with col2:
        # Leadership radar chart
        leadership_chart = create_leadership_influence_chart(analysis_360)
        if leadership_chart:
            st.plotly_chart(leadership_chart, use_container_width=True)

def display_growth_analysis(analysis_360):
    """Display growth and development analysis"""
    growth_data = analysis_360['growth_trajectory']
    cross_func_data = analysis_360['cross_functional_contribution']
    innovation_data = analysis_360['innovation_contribution']
    performance_data = analysis_360['performance_comparison']
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📈 Xu hướng Phát triển")
        st.markdown(f"**Xu hướng tăng trưởng:** {growth_data['growth_trend']}")
        st.markdown(f"**Điểm quỹ đạo:** {growth_data['trajectory_score']:.1f}/10")
        
        if 'completion_rate_change' in growth_data:
            st.markdown(f"**Thay đổi tỷ lệ hoàn thành:** {growth_data['completion_rate_change']:.1f}%")
        
        if 'project_count_change' in growth_data:
            st.markdown(f"**Thay đổi số dự án:** {growth_data['project_count_change']}")
        
        st.markdown("### 🔄 Khả năng Đa chức năng")
        st.markdown(f"**Điểm đa chức năng:** {cross_func_data['cross_functional_score']:.1f}/10")
        st.markdown(f"**Loại kỹ năng:** {cross_func_data['skill_categories']}")
        st.markdown(f"**Loại dự án:** {cross_func_data['project_types']}")
        st.markdown(f"**Mức độ linh hoạt:** {cross_func_data['versatility_level']}")
    
    with col2:
        st.markdown("### 💡 Đổi mới & Sáng tạo")
        st.markdown(f"**Điểm đổi mới:** {innovation_data['innovation_score']:.1f}/10")
        st.markdown(f"**Đa dạng dự án:** {innovation_data['project_diversity']}")
        st.markdown(f"**Bề rộng kỹ năng:** {innovation_data['skill_breadth']}")
        st.markdown(f"**Mức độ đổi mới:** {innovation_data['innovation_level']}")
        
        st.markdown("### 🏆 So sánh với Đồng nghiệp")
        if isinstance(performance_data, dict) and 'peer_count' in performance_data:
            st.markdown(f"**Số đồng nghiệp so sánh:** {performance_data['peer_count']}")
            st.markdown(f"**Percentile hoàn thành:** {performance_data['completion_rate_percentile']}%")
            st.markdown(f"**Percentile dự án:** {performance_data['project_count_percentile']}%")
            st.markdown(f"**Hiệu suất tương đối:** {performance_data['relative_performance']}")
        else:
            st.markdown(f"**Kết quả so sánh:** {performance_data}")

def display_risk_recommendations(analysis_360):
    """Display risk assessment and recommendations"""
    risk_data = analysis_360['risk_assessment']
    recommendations = analysis_360['recommendations_360']
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### ⚠️ Đánh giá Rủi ro")
        st.markdown(f"**Mức độ rủi ro:** {risk_data['risk_level']}")
        st.markdown(f"**Điểm rủi ro:** {risk_data['risk_score']}/10")
        
        if risk_data['identified_risks']:
            st.markdown("**Rủi ro được xác định:**")
            for risk in risk_data['identified_risks']:
                st.markdown(f"• {risk}")
        else:
            st.markdown("✅ Không có rủi ro đáng kể được xác định")
    
    with col2:
        st.markdown("### 💡 Khuyến nghị 360")
        if recommendations:
            for i, rec in enumerate(recommendations, 1):
                st.markdown(f"{i}. {rec}")
        else:
            st.markdown("✅ Hiệu suất tốt, tiếp tục duy trì")
    
    # Action plan
    st.markdown("### 📋 Kế hoạch Hành động Đề xuất")
    
    action_plan = []
    
    # Based on networking score
    networking_score = analysis_360['collaboration_analysis']['networking_score']
    if networking_score < 5:
        action_plan.append("🤝 Tăng cường hoạt động networking và hợp tác nhóm")
    elif networking_score > 8:
        action_plan.append("🌟 Chia sẻ kinh nghiệm networking với đồng nghiệp")
    
    # Based on leadership score
    leadership_score = analysis_360['leadership_influence']['leadership_score']
    if leadership_score < 5:
        action_plan.append("👑 Tham gia các khóa đào tạo lãnh đạo")
    elif leadership_score > 7:
        action_plan.append("🎯 Đảm nhận vai trò mentor cho nhân viên mới")
    
    # Based on innovation score
    innovation_score = analysis_360['innovation_contribution']['innovation_score']
    if innovation_score < 6:
        action_plan.append("💡 Tham gia các dự án sáng tạo và đổi mới")
    
    # Based on risk level
    if risk_data['risk_score'] > 5:
        action_plan.append("⚠️ Ưu tiên giải quyết các rủi ro được xác định")
    
    if action_plan:
        for i, action in enumerate(action_plan, 1):
            st.markdown(f"{i}. {action}")
    else:
        st.markdown("✅ Hiệu suất toàn diện tốt, tiếp tục phát triển theo hướng hiện tại")

def export_360_to_word(analysis_360, employee_info, time_filter):
    """Export Employee 360 analysis to Word document"""
    doc = Document()
    
    # Title
    title = doc.add_heading(f'Báo cáo Phân tích 360: {employee_info["name"]}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(f'Khoảng thời gian: {time_filter} | Phân tích toàn diện 360 độ')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Employee info
    doc.add_heading('👤 Thông tin Nhân viên', level=1)
    p = doc.add_paragraph()
    p.add_run('Tên: ').bold = True
    p.add_run(employee_info['name'])
    p.add_run(' | Username: ').bold = True
    p.add_run(employee_info['username'])
    p.add_run(' | Vị trí: ').bold = True
    p.add_run(employee_info['job'])
    
    # Overview scores
    doc.add_heading('📊 Tổng quan Điểm số 360', level=1)
    scores_table = doc.add_table(rows=7, cols=2)
    scores_table.style = 'Table Grid'
    
    scores_data = [
        ('Networking', f"{analysis_360['collaboration_analysis']['networking_score']:.1f}/10"),
        ('Chuyên môn', f"{analysis_360['skill_analysis']['expertise_level']:.1f}/10"),
        ('Lãnh đạo', f"{analysis_360['leadership_influence']['leadership_score']:.1f}/10"),
        ('Đa chức năng', f"{analysis_360['cross_functional_contribution']['cross_functional_score']:.1f}/10"),
        ('Tác động nhóm', f"{analysis_360['team_impact']['team_impact_score']:.1f}/10"),
        ('Đổi mới', f"{analysis_360['innovation_contribution']['innovation_score']:.1f}/10")
    ]
    
    for i, (metric, score) in enumerate(scores_data):
        scores_table.cell(i, 0).text = metric
        scores_table.cell(i, 1).text = score
    
    # Detailed sections
    sections = [
        ('🤝 Hợp tác & Mạng lưới', analysis_360['collaboration_analysis']),
        ('🎯 Kỹ năng & Chuyên môn', analysis_360['skill_analysis']),
        ('👑 Lãnh đạo & Ảnh hưởng', analysis_360['leadership_influence']),
        ('🔄 Khả năng Đa chức năng', analysis_360['cross_functional_contribution']),
        ('📈 Phát triển & Xu hướng', analysis_360['growth_trajectory']),
        ('💡 Đổi mới & Sáng tạo', analysis_360['innovation_contribution'])
    ]
    
    for section_title, section_data in sections:
        doc.add_heading(section_title, level=1)
        
        # Add key metrics for each section
        for key, value in section_data.items():
            if isinstance(value, (int, float)):
                p = doc.add_paragraph()
                p.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                p.add_run(str(value))
            elif isinstance(value, str) and len(value) < 100:
                p = doc.add_paragraph()
                p.add_run(f'{key.replace("_", " ").title()}: ').bold = True
                p.add_run(value)
    
    # Risk assessment
    doc.add_heading('⚠️ Đánh giá Rủi ro', level=1)
    risk_data = analysis_360['risk_assessment']
    p = doc.add_paragraph()
    p.add_run('Mức độ rủi ro: ').bold = True
    p.add_run(risk_data['risk_level'])
    
    if risk_data['identified_risks']:
        doc.add_heading('Rủi ro được xác định:', level=2)
        for risk in risk_data['identified_risks']:
            doc.add_paragraph(f'• {risk}')
    
    # Recommendations
    doc.add_heading('💡 Khuyến nghị 360', level=1)
    recommendations = analysis_360['recommendations_360']
    for i, rec in enumerate(recommendations, 1):
        doc.add_paragraph(f'{i}. {rec}')
    
    return doc

def employee_analysis_interface(api_client, basic_analyzer, ai_analyzer, analysis_mode):
    """Interface for employee time-based analysis"""
    
    st.header("👥 Phân tích Nhân viên theo Thời gian")
    
    # Load employees
    if 'employees' not in st.session_state:
        with st.spinner("Đang tải danh sách nhân viên..."):
            try:
                st.session_state.employees = api_client.get_filtered_members()
            except Exception as e:
                st.error(f"❌ Lỗi khi tải danh sách nhân viên: {e}")
                return
    
    employees_df = st.session_state.employees
    
    if employees_df.empty:
        st.error("❌ Không có nhân viên nào được tìm thấy")
        return
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Employee selection
        st.subheader("📋 Chọn Nhân viên")
        
        employee_options = {}
        for _, employee in employees_df.iterrows():
            display_name = f"{employee['name']} ({employee['username']}) - {employee['job']}"
            employee_options[display_name] = employee.to_dict()
        
        selected_employee_display = st.selectbox(
            "Chọn nhân viên để phân tích:",
            options=list(employee_options.keys()),
            help="Chọn nhân viên từ danh sách"
        )
        
        # Time period selection
        st.subheader("📅 Chọn Khoảng thời gian")
        
        time_filter = st.selectbox(
            "Khoảng thời gian phân tích:",
            [
                "Tháng hiện tại",
                "Quý hiện tại", 
                "Năm hiện tại",
                "3 tháng gần nhất",
                "6 tháng gần nhất",
                "Tất cả"
            ],
            index=1,  # Default to current quarter
            help="Chọn khoảng thời gian để phân tích hiệu suất"
        )
    
    with col2:
        st.subheader("📊 Thông tin Nhân viên")
        
        if selected_employee_display:
            selected_employee = employee_options[selected_employee_display]
            
            st.markdown(f"**Tên:** {selected_employee['name']}")
            st.markdown(f"**Username:** {selected_employee['username']}")
            st.markdown(f"**Vị trí:** {selected_employee['job']}")
            st.markdown(f"**Email:** {selected_employee['email']}")
            
            st.markdown(f"**Phân tích:** {time_filter}")
            st.markdown(f"**Phương pháp:** {analysis_mode}")
            
            # AI model info
            if ai_analyzer and ai_analyzer.client and analysis_mode in ["AI", "Cả hai"]:
                model_name = "Gemma" if ai_analyzer.use_gemma else "Gemini"
                st.info(f"🤖 Sử dụng {model_name} AI")
            
            # Analysis button
            if st.button("🔄 Phân tích Nhân viên", type="primary"):
                username = selected_employee['username']
                
                with st.spinner(f"Đang thu thập dữ liệu {time_filter.lower()}..."):
                    employee_projects = api_client.get_employee_projects_by_time(username, time_filter)
                
                if employee_projects:
                    # Parse all tasks into DataFrame
                    analyzer = TaskAnalyzer()
                    df = analyzer.parse_employee_tasks(employee_projects)
                    
                    if not df.empty:
                        st.session_state.employee_analysis_data = {
                            'df': df,
                            'employee_info': selected_employee,
                            'time_filter': time_filter,
                            'projects_data': employee_projects,
                            'analysis_mode': analysis_mode
                        }
                        
                        # Basic Analysis
                        if analysis_mode in ["Cơ bản", "Cả hai"]:
                            with st.spinner("📊 Đang thực hiện phân tích cơ bản..."):
                                basic_analysis = basic_analyzer.analyze_employee_performance_basic(
                                    employee_projects,
                                    selected_employee,
                                    time_filter
                                )
                                st.session_state.employee_analysis_data['basic_analysis'] = basic_analysis
                        
                        # AI Analysis
                        if ai_analyzer and ai_analyzer.client and analysis_mode in ["AI", "Cả hai"]:
                            model_name = "Gemma" if ai_analyzer.use_gemma else "Gemini"
                            with st.spinner(f"🤖 Đang phân tích với {model_name} AI..."):
                                ai_analysis = ai_analyzer.analyze_employee_performance_by_time(
                                    employee_projects,
                                    selected_employee,
                                    time_filter
                                )
                                
                                if ai_analysis:
                                    st.session_state.employee_analysis_data['ai_analysis'] = ai_analysis
                        
                        st.success("✅ Phân tích thành công!")
                    else:
                        st.warning(f"⚠️ Không tìm thấy công việc nào trong {time_filter.lower()}")
                else:
                    st.warning(f"⚠️ Không tìm thấy dữ liệu cho {selected_employee['name']} trong {time_filter.lower()}")
    
    # Display results
    display_employee_analysis_results()

def display_employee_analysis_results():
    """Display employee analysis results"""
    if 'employee_analysis_data' not in st.session_state:
        return
    
    data = st.session_state.employee_analysis_data
    df = data['df']
    employee_info = data['employee_info']
    time_filter = data['time_filter']
    analysis_mode = data.get('analysis_mode', 'Cơ bản')
    
    st.markdown("---")
    st.header(f"📊 Kết quả Phân tích: {employee_info['name']} - {time_filter}")
    
    # Create tabs based on analysis mode
    if analysis_mode == "Cả hai":
        tab1, tab2 = st.tabs(["📊 Phân tích Cơ bản", "🤖 Phân tích AI"])
        
        with tab1:
            display_basic_analysis_results(data)
        
        with tab2:
            display_ai_analysis_results(data)
    
    elif analysis_mode == "AI":
        display_ai_analysis_results(data)
    
    else:  # Basic analysis
        display_basic_analysis_results(data)
    
    # Common sections
    display_common_analysis_sections(data)

def display_basic_analysis_results(data):
    """Display basic analysis results"""
    if 'basic_analysis' not in data:
        st.error("❌ Không có kết quả phân tích cơ bản")
        return
    
    basic_result = data['basic_analysis']
    stats = basic_result['overall_stats']
    data_quality = basic_result['data_quality']
    
    st.subheader("📊 Phân tích Cơ bản")
    
    # Statistics including ongoing tasks and on-time completion rate
    col1, col2, col3, col4, col5 = st.columns(5)
    
    with col1:
        st.metric("Tổng dự án", stats['total_projects'])
    
    with col2:
        st.metric("Tổng công việc", stats['total_tasks'])
    
    with col3:
        st.metric("Hoàn thành", stats['completed_tasks'])
    
    with col4:
        st.metric("Đang thực hiện", stats['ongoing_tasks'])
    
    with col5:
        st.metric("Hoàn thành đúng hạn", f"{stats['on_time_completion_rate']}%")
    
    # Performance chart
    col1, col2 = st.columns(2)
    
    with col1:
        performance_fig = create_basic_performance_chart(basic_result)
        st.plotly_chart(performance_fig, use_container_width=True)
    
    with col2:
        # Data quality
        st.subheader("🎯 Chất lượng Dữ liệu")
        st.metric("Điểm chất lượng", f"{data_quality['average_score']}/100")
        st.metric("Xếp hạng", data_quality['grade'])
        
        if data_quality['common_issues']:
            st.write("**Vấn đề chung:**")
            for issue in data_quality['common_issues']:
                st.write(f"• {issue}")
    
    # Data quality by project
    quality_fig = create_data_quality_chart(basic_result)
    if quality_fig:
        st.plotly_chart(quality_fig, use_container_width=True)
    
    # Project analysis details
    st.subheader("📋 Phân tích Chi tiết theo Dự án")
    
    for project_analysis in basic_result['projects_analysis']:
        with st.expander(f"📁 {project_analysis['project_name']} - Điểm: {project_analysis['quality_score']}/100"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.write("**Thống kê:**")
                st.write(f"• Tổng task: {project_analysis['total_tasks']}")
                st.write(f"• Điểm chất lượng: {project_analysis['quality_score']}/100")
            
            with col2:
                st.write("**Trường quan trọng bị thiếu:**")
                important_missing_fields = project_analysis.get('important_missing_fields', {})
                if important_missing_fields:
                    for field, count in important_missing_fields.items():
                        if count > 0:
                            st.write(f"• {field}: {count} task")
                else:
                    st.write("• Không có trường quan trọng nào bị thiếu")
            
            with col3:
                st.write("**Vấn đề:**")
                if project_analysis['issues']:
                    for issue in project_analysis['issues']:
                        st.write(f"• {issue}")
                else:
                    st.write("• Không có vấn đề đặc biệt")
    
    # Recommendations
    st.subheader("💡 Khuyến nghị")
    for rec in basic_result['recommendations']:
        st.write(f"• {rec}")

def display_ai_analysis_results(data):
    """Display AI analysis results"""
    if 'ai_analysis' not in data:
        st.error("❌ Không có kết quả phân tích AI")
        return
    
    ai_result = data['ai_analysis']
    
    st.subheader("🤖 Đánh giá AI")
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        # Performance metrics
        col_a, col_b, col_c = st.columns(3)
        with col_a:
            st.metric("Hiệu suất tổng thể", f"{ai_result.get('hieu_suat_tong_the', 0)}/10")
        with col_b:
            st.metric("Khả năng hoàn thành", f"{ai_result.get('kha_nang_hoan_thanh', 0)}/10")
        with col_c:
            st.metric("Chất lượng công việc", f"{ai_result.get('chat_luong_cong_viec', 0)}/10")
        
        col_d, col_e, col_f = st.columns(3)
        with col_d:
            st.metric("Tuân thủ deadline", f"{ai_result.get('tuan_thu_deadline', 0)}/10")
        with col_e:
            st.metric("Tính nhất quán", f"{ai_result.get('tinh_nhat_quan', 0)}/10")
        with col_f:
            st.metric("Khả năng đa dự án", f"{ai_result.get('kha_nang_da_du_an', 0)}/10")
    
    with col2:
        # Performance radar chart
        radar_fig = create_employee_performance_radar(ai_result)
        st.plotly_chart(radar_fig, use_container_width=True)
    
    # Detailed AI analysis
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.write("**🎯 Điểm mạnh:**")
        for strength in ai_result.get('diem_manh', []):
            st.write(f"• {strength}")
    
    with col2:
        st.write("**⚠️ Điểm cần cải thiện:**")
        for improvement in ai_result.get('diem_can_cai_thien', []):
            st.write(f"• {improvement}")
    
    with col3:
        st.write("**💡 Khuyến nghị phát triển:**")
        for recommendation in ai_result.get('khuyen_nghi_phat_trien', []):
            st.write(f"• {recommendation}")
    
    # Trend analysis
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📈 Xu hướng Hiệu suất")
        st.write(ai_result.get('xu_huong_hieu_suat', 'Không có đánh giá'))
        
        if ai_result.get('so_sanh_ky_truoc'):
            st.subheader("🔄 So sánh với kỳ trước")
            st.write(ai_result.get('so_sanh_ky_truoc'))
    
    with col2:
        st.subheader("🎯 Khuyến nghị kỳ tiếp theo")
        st.write(ai_result.get('khuyen_nghi_ky_tiep', 'Không có khuyến nghị'))

def display_common_analysis_sections(data):
    """Display common analysis sections"""
    df = data['df']
    employee_info = data['employee_info']
    time_filter = data['time_filter']
    analysis_mode = data.get('analysis_mode', 'Cơ bản')
    
    # Visualizations
    st.subheader("📊 Biểu đồ Phân tích")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Project distribution
        project_fig = create_project_distribution_chart(df)
        if project_fig:
            st.plotly_chart(project_fig, use_container_width=True)
    
    with col2:
        # Status distribution
        if 'Trạng thái' in df.columns:
            status_counts = df['Trạng thái'].value_counts()
            fig_pie = px.pie(values=status_counts.values, names=status_counts.index, 
                           title="Phân bố Trạng thái")
            st.plotly_chart(fig_pie, use_container_width=True)
    
    # Timeline chart
    timeline_fig = create_timeline_chart(df)
    if timeline_fig:
        st.plotly_chart(timeline_fig, use_container_width=True)
    
    # Data table
    st.subheader("📋 Chi tiết Công việc")
    st.dataframe(df, use_container_width=True)
    
    # Export section
    st.markdown("---")
    st.header("📤 Xuất Báo cáo")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📄 Tải báo cáo Word", type="primary"):
            try:
                analyzer = TaskAnalyzer()
                
                # Determine analysis result for Word export
                analysis_result = None
                export_analysis_type = "Cơ bản"
                
                if analysis_mode == "AI" and 'ai_analysis' in data:
                    analysis_result = {'ai_analysis': data['ai_analysis']}
                    export_analysis_type = "AI"
                elif analysis_mode == "Cơ bản" and 'basic_analysis' in data:
                    analysis_result = data['basic_analysis']
                    export_analysis_type = "Cơ bản"
                elif analysis_mode == "Cả hai":
                    # For combined mode, export both if available
                    analysis_result = {}
                    if 'basic_analysis' in data:
                        analysis_result.update(data['basic_analysis'])
                    if 'ai_analysis' in data:
                        analysis_result['ai_analysis'] = data['ai_analysis']
                    export_analysis_type = "Kết hợp"
                
                doc = analyzer.export_employee_report_to_word(
                    df, 
                    employee_info['name'], 
                    time_filter,
                    analysis_result,
                    export_analysis_type
                )
                
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                st.download_button(
                    label="📥 Tải xuống file Word",
                    data=doc_buffer.getvalue(),
                    file_name=f"{employee_info['username']}_{time_filter}_{analysis_mode}_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                st.success("✅ Báo cáo Word đã được tạo!")
                
            except Exception as e:
                st.error(f"❌ Lỗi khi tạo báo cáo Word: {str(e)}")
    
    with col2:
        csv_data = df.to_csv(index=False, encoding='utf-8-sig')
        st.download_button(
            label="📊 Tải dữ liệu CSV",
            data=csv_data,
            file_name=f"{employee_info['username']}_{time_filter}_data.csv",
            mime="text/csv"
        )
    
    with col3:
        # Export analysis results
        export_data = {
            'employee_info': employee_info,
            'time_filter': time_filter,
            'analysis_mode': analysis_mode,
            'statistics': {
                'total_tasks': len(df),
                'completed_tasks': len(df[df['Trạng thái'] == 'Hoàn thành']) if 'Trạng thái' in df.columns else 0,
                'ongoing_tasks': len(df[df['Trạng thái'] == 'Đang thực hiện']) if 'Trạng thái' in df.columns else 0,
                'projects_count': len(df['Dự án'].unique()) if 'Dự án' in df.columns else 0
            }
        }
        
        if 'basic_analysis' in data:
            export_data['basic_analysis'] = data['basic_analysis']
        
        if 'ai_analysis' in data:
            export_data['ai_analysis'] = data['ai_analysis']
        
        json_data = json.dumps(export_data, ensure_ascii=False, indent=2)
        st.download_button(
            label="📋 Tải kết quả phân tích",
            data=json_data,
            file_name=f"{employee_info['username']}_{time_filter}_{analysis_mode}_analysis.json",
            mime="application/json"
        )

if __name__ == "__main__":
    main()
